"""
Chamorro Chatbot FastAPI Application

A simple API wrapper around the chatbot core logic.
"""

from fastapi import FastAPI, HTTPException, Request, Header, File, UploadFile, Form, BackgroundTasks, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
import time
import os
import logging
import base64
import queue
import threading
import asyncio
import boto3
from botocore.exceptions import ClientError
from dotenv import load_dotenv
from collections import defaultdict
from datetime import datetime, timedelta
from typing import Optional, List
import json

from .models import (
    ChatRequest,
    ChatResponse,
    HealthResponse,
    ErrorResponse,
    SourceInfo,
    ConversationCreate,
    ConversationResponse,
    ConversationListResponse,
    MessagesResponse,
    SystemMessageCreate,
    InitResponse,
    FlashcardResponse,
    GenerateFlashcardsResponse,
    # Flashcard Progress Models
    SaveDeckRequest,
    SaveDeckResponse,
    FlashcardProgressInfo,
    FlashcardWithProgress,
    UserDeckResponse,
    UserDecksResponse,
    DeckCardsResponse,
    ReviewCardRequest,
    ReviewCardResponse,
    # Feedback Models
    MessageFeedbackRequest,
    MessageFeedbackResponse,
    # Quiz Result Models
    QuizResultCreate,
    QuizResultResponse,
    QuizResultDetailResponse,
    QuizAnswerResponse,
    QuizStatsResponse,
    # Game Result Models
    GameResultCreate,
    GameResultResponse,
    GameStatsResponse,
    # Admin Dashboard Models
    AdminStatsResponse,
    AdminUserInfo,
    AdminUsersResponse,
    AdminUserUpdateRequest,
    AdminUserUpdateResponse,
    # Share Conversation Models
    ShareConversationRequest,
    ShareConversationResponse,
    SharedConversationResponse,
    ShareInfoResponse
)
from .chatbot_service import get_chatbot_response, get_chatbot_response_stream, cancel_pending_message
from . import conversations

# Clerk for authentication
try:
    from clerk_backend_api import Clerk
    CLERK_AVAILABLE = True
except ImportError:
    CLERK_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("⚠️  clerk-backend-api not installed. Authentication disabled.")

# Add parent directory to path for root-level imports
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

# DON'T import heavy modules at startup - lazy load them!
# This saves ~400MB of memory on free tier
# from src.rag.chamorro_rag import rag
# RAGDatabaseManager is only needed for health check stats, not API runtime
# from src.rag.manage_rag_db import RAGDatabaseManager

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Sentry for error tracking (if configured)
try:
    from src.utils.sentry_config import init_sentry, set_user_context, set_request_context
    sentry_enabled = init_sentry()
except ImportError as e:
    logger.warning(f"⚠️  Sentry not available: {e}")
    sentry_enabled = False
    # Create dummy functions if Sentry not available
    def set_user_context(*args, **kwargs): pass
    def set_request_context(*args, **kwargs): pass

# Initialize Clerk client (if available and configured)
clerk = None
if CLERK_AVAILABLE and os.getenv("CLERK_SECRET_KEY"):
    try:
        clerk = Clerk(bearer_auth=os.getenv("CLERK_SECRET_KEY"))
        logger.info(f"✅ Clerk initialized: {os.getenv('CLERK_SECRET_KEY')[:15]}...")
    except Exception as e:
        logger.error(f"❌ Failed to initialize Clerk: {e}")
else:
    if not os.getenv("CLERK_SECRET_KEY"):
        logger.warning("⚠️  CLERK_SECRET_KEY not set. Running without authentication.")

async def verify_user(authorization: Optional[str] = Header(None)) -> str:
    """
    Verify Clerk JWT token and return user ID.
    
    Authentication is REQUIRED - raises HTTPException if not authenticated.
    
    Args:
        authorization: Bearer token from request header
        
    Returns:
        user_id: Clerk user ID
        
    Raises:
        HTTPException: 401 if not authenticated or invalid token
    """
    from jose import jwt
    from jose.exceptions import JWTError
    
    if not authorization:
        raise HTTPException(
            status_code=401,
            detail="Authentication required. Please sign in to use HåfaGPT."
        )
    
    if not clerk:
        logger.error("⚠️  Clerk not initialized but authentication required")
        raise HTTPException(
            status_code=500,
            detail="Authentication service unavailable"
        )
    
    try:
        # Extract token from "Bearer <token>"
        token = authorization.replace("Bearer ", "")
        
        # Get JWKS from Clerk - use get_jwks() method
        jwks_response = clerk.jwks.get_jwks()
        
        # Decode and verify JWT
        # First, decode without verification to get the kid (key ID)
        unverified_header = jwt.get_unverified_header(token)
        kid = unverified_header.get('kid')
        
        # Find the matching key in JWKS
        jwks_keys = jwks_response.keys if hasattr(jwks_response, 'keys') else []
        matching_key = None
        for key in jwks_keys:
            if hasattr(key, 'kid') and key.kid == kid:
                matching_key = key
                break
        
        if not matching_key:
            logger.warning(f"⚠️  No matching key found for kid: {kid}")
            raise HTTPException(
                status_code=401,
                detail="Invalid authentication token. Please sign in again."
            )
        
        # Verify and decode the JWT
        # Convert the key to dict format for jose
        key_dict = {
            'kty': matching_key.kty,
            'use': matching_key.use,
            'kid': matching_key.kid,
            'n': matching_key.n,
            'e': matching_key.e,
        }
        
        # Decode the token
        payload = jwt.decode(
            token,
            key_dict,
            algorithms=['RS256'],
            options={'verify_aud': False}  # Clerk doesn't always set aud
        )
        
        user_id = payload.get('sub')
        
        if not user_id:
            logger.warning("⚠️  JWT verified but no user_id in payload")
            raise HTTPException(
                status_code=401,
                detail="Invalid user data in token. Please sign in again."
            )
        
        logger.info(f"✅ Authenticated user: {user_id}")
        return user_id
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except JWTError as e:
        logger.warning(f"⚠️  JWT verification failed: {e}")
        raise HTTPException(
            status_code=401,
            detail="Invalid or expired token. Please sign in again."
        )
    except Exception as e:
        logger.warning(f"⚠️  Authentication error: {e}")
        raise HTTPException(
            status_code=401,
            detail="Authentication failed. Please sign in again."
        )

# Initialize S3 client for image uploads
try:
    s3_client = boto3.client(
        's3',
        aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
        region_name=os.getenv('AWS_REGION', 'us-west-2')
    )
    S3_BUCKET = os.getenv('AWS_S3_BUCKET')
    S3_AVAILABLE = bool(S3_BUCKET)
    if S3_AVAILABLE:
        logger.info(f"✅ S3 client initialized for bucket: {S3_BUCKET}")
    else:
        logger.warning("⚠️  AWS_S3_BUCKET not set. Image uploads will not be persisted.")
except Exception as e:
    S3_AVAILABLE = False
    logger.warning(f"⚠️  Failed to initialize S3 client: {e}")

def upload_image_to_s3(image_data: bytes, filename: str, content_type: str) -> Optional[str]:
    """
    Upload image to S3 and return public URL.
    
    Args:
        image_data: Binary image data
        filename: Original filename
        content_type: MIME type (e.g., 'image/jpeg')
    
    Returns:
        S3 public URL or None if upload fails
    """
    if not S3_AVAILABLE:
        logger.warning("S3 not available, skipping image upload")
        return None
    
    try:
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_extension = filename.split('.')[-1] if '.' in filename else 'jpg'
        s3_key = f"chamorro_uploads/{timestamp}_{filename}"
        
        # Upload to S3 (without ACL - bucket policy handles public access)
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=image_data,
            ContentType=content_type
            # No ACL needed - bucket policy makes objects public
        )
        
        # Construct public URL
        image_url = f"https://{S3_BUCKET}.s3.{os.getenv('AWS_REGION', 'us-west-2')}.amazonaws.com/{s3_key}"
        logger.info(f"✅ Image uploaded to S3: {image_url}")
        return image_url
        
    except ClientError as e:
        logger.error(f"Failed to upload image to S3: {e}")
        return None
    except Exception as e:
        logger.error(f"Unexpected error uploading to S3: {e}")
        return None


# --- File Upload Support (PDF, Word, Text) ---

import io
from pypdf import PdfReader
from docx import Document

# Supported file types and their MIME types
SUPPORTED_FILE_TYPES = {
    # Documents
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',  # Old .doc format (limited support)
    'text/plain': 'txt',
    # Images (existing support)
    'image/jpeg': 'image',
    'image/png': 'image',
    'image/webp': 'image',
    'image/gif': 'image',
}

def extract_text_from_pdf(file_data: bytes) -> str:
    """Extract text from PDF file."""
    try:
        pdf = PdfReader(io.BytesIO(file_data))
        text_parts = []
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"📄 PDF extracted: {len(pdf.pages)} pages, {len(full_text)} chars")
        return full_text
    except Exception as e:
        logger.error(f"❌ Failed to extract PDF text: {e}")
        raise ValueError(f"Failed to read PDF: {str(e)}")


def extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from Word document (.docx)."""
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"📝 Word doc extracted: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars")
        return full_text
    except Exception as e:
        logger.error(f"❌ Failed to extract Word doc text: {e}")
        raise ValueError(f"Failed to read Word document: {str(e)}")


def extract_text_from_txt(file_data: bytes) -> str:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            text = file_data.decode('utf-8')
        except UnicodeDecodeError:
            text = file_data.decode('latin-1')
        
        logger.info(f"📋 Text file extracted: {len(text)} chars")
        return text
    except Exception as e:
        logger.error(f"❌ Failed to read text file: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")


def upload_file_to_s3(file_data: bytes, filename: str, content_type: str) -> Optional[str]:
    """
    Upload file to S3 and return public URL.
    Works for both images and documents.
    """
    if not S3_AVAILABLE:
        logger.warning("S3 not available, skipping file upload")
        return None
    
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_extension = filename.split('.')[-1] if '.' in filename else 'bin'
        s3_key = f"chamorro_uploads/{timestamp}_{filename}"
        
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file_data,
            ContentType=content_type
        )
        
        file_url = f"https://{S3_BUCKET}.s3.{os.getenv('AWS_REGION', 'us-west-2')}.amazonaws.com/{s3_key}"
        logger.info(f"✅ File uploaded to S3: {file_url}")
        return file_url
        
    except ClientError as e:
        logger.error(f"Failed to upload file to S3: {e}")
        return None
    except Exception as e:
        logger.error(f"Unexpected error uploading to S3: {e}")
        return None


def process_uploaded_file(file_data: bytes, content_type: str, filename: str) -> dict:
    """
    Process an uploaded file and return extracted content.
    
    Returns:
        dict with keys:
        - file_type: 'pdf', 'docx', 'txt', or 'image'
        - text_content: Extracted text (for documents) or None (for images)
        - image_base64: Base64 encoded image (for images) or None (for documents)
    """
    file_type = SUPPORTED_FILE_TYPES.get(content_type)
    
    if not file_type:
        raise ValueError(f"Unsupported file type: {content_type}. Supported: PDF, Word (.docx), Text (.txt), Images")
    
    result = {
        'file_type': file_type,
        'text_content': None,
        'image_base64': None,
        'filename': filename
    }
    
    if file_type == 'pdf':
        result['text_content'] = extract_text_from_pdf(file_data)
    elif file_type == 'docx':
        result['text_content'] = extract_text_from_docx(file_data)
    elif file_type == 'doc':
        # Old .doc format - limited support, suggest converting
        raise ValueError("Old .doc format not supported. Please save as .docx (Word 2007+)")
    elif file_type == 'txt':
        result['text_content'] = extract_text_from_txt(file_data)
    elif file_type == 'image':
        result['image_base64'] = base64.b64encode(file_data).decode('utf-8')
    
    return result


# ============================================================================
# Background S3 Upload Functions
# ============================================================================

def append_file_url_to_conversation_log(conversation_id: str, file_info: dict):
    """
    Append a file URL to the conversation log's file_urls array.
    Called from background task after each S3 upload completes.
    
    Args:
        conversation_id: The conversation ID
        file_info: Dict with {url, filename, type} - type is 'image' or 'document'
    """
    try:
        import psycopg
        import json
        conn = psycopg.connect(os.getenv("DATABASE_URL"))
        cursor = conn.cursor()
        
        # Append to file_urls JSON array (create array if null)
        # Also update legacy image_url field for backwards compatibility
        cursor.execute("""
            UPDATE conversation_logs 
            SET 
                file_urls = COALESCE(file_urls, '[]'::jsonb) || %s::jsonb,
                image_url = COALESCE(image_url, %s)
            WHERE id = (
                SELECT id FROM conversation_logs 
                WHERE conversation_id = %s
                ORDER BY timestamp DESC
                LIMIT 1
            )
        """, (json.dumps([file_info]), file_info['url'], conversation_id))
        
        rows_updated = cursor.rowcount
        conn.commit()
        cursor.close()
        conn.close()
        
        if rows_updated > 0:
            logger.info(f"✅ Background: Added {file_info['filename']} to file_urls")
        else:
            logger.debug(f"No conversation_log to update for {conversation_id}")
    except Exception as e:
        logger.error(f"Failed to append file_url: {e}")


def upload_file_to_s3_background(
    file_data: bytes,
    filename: str,
    content_type: str,
    user_id: str,
    conversation_id: str,
    file_index: int
):
    """
    Background task to upload file to S3 and update conversation_logs.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        content_type: MIME type
        user_id: User ID for S3 path
        conversation_id: Conversation ID for DB update
        file_index: Index for unique filename
    """
    try:
        if not s3_client:
            logger.warning("⚠️ S3 client not configured, skipping background upload")
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        s3_key = f"uploads/{user_id or 'anonymous'}/{timestamp}_{file_index}_{filename}"
        
        # Upload to S3
        logger.info(f"📤 Background: Uploading {filename} to S3...")
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file_data,
            ContentType=content_type
        )
        
        file_url = f"https://{S3_BUCKET}.s3.{os.getenv('AWS_REGION', 'us-west-2')}.amazonaws.com/{s3_key}"
        logger.info(f"✅ Background: Uploaded {filename} → {file_url}")
        
        # Determine file type for display
        file_type = 'image' if content_type.startswith('image/') else 'document'
        
        # Append to conversation_logs file_urls array
        if conversation_id:
            append_file_url_to_conversation_log(conversation_id, {
                'url': file_url,
                'filename': filename,
                'type': file_type,
                'content_type': content_type
            })
            
    except Exception as e:
        logger.error(f"❌ Background S3 upload error for {filename}: {e}")


# Rate limiting storage (in-memory - use Redis in production for multiple servers)
rate_limit_storage = defaultdict(list)
RATE_LIMIT_REQUESTS = int(os.getenv("RATE_LIMIT_REQUESTS", "60"))  # requests
RATE_LIMIT_WINDOW = int(os.getenv("RATE_LIMIT_WINDOW", "60"))  # seconds

def check_rate_limit(client_ip: str) -> bool:
    """
    Check if client has exceeded rate limit.
    Returns True if allowed, False if rate limited.
    """
    now = datetime.now()
    cutoff = now - timedelta(seconds=RATE_LIMIT_WINDOW)
    
    # Clean old requests
    rate_limit_storage[client_ip] = [
        req_time for req_time in rate_limit_storage[client_ip]
        if req_time > cutoff
    ]
    
    # Check limit
    if len(rate_limit_storage[client_ip]) >= RATE_LIMIT_REQUESTS:
        return False
    
    # Add current request
    rate_limit_storage[client_ip].append(now)
    return True

# Create FastAPI app
app = FastAPI(
    title="Chamorro Language Learning Chatbot API",
    description="AI-powered Chamorro language tutor with RAG and web search",
    version="1.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# Log startup
logger.info("🚀 FastAPI app created successfully")
logger.info(f"📍 PORT environment variable: {os.getenv('PORT', 'NOT SET')}")

# Get allowed origins from environment variable
# Format: "https://domain1.com,https://domain2.com" or "*" for development
allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "*")
allowed_origins = [origin.strip() for origin in allowed_origins_str.split(",")] if allowed_origins_str != "*" else ["*"]

# Always allow Capacitor iOS app origins
CAPACITOR_ORIGINS = [
    "capacitor://localhost",  # iOS Capacitor app
    "http://localhost",       # Android Capacitor app
    "ionic://localhost",      # Ionic apps
]
# Add Capacitor origins if not already using wildcard
if allowed_origins != ["*"]:
    allowed_origins.extend(CAPACITOR_ORIGINS)

# Log CORS configuration
if allowed_origins == ["*"]:
    logger.warning("⚠️  CORS is set to allow all origins (*). This is OK for development but should be restricted in production!")
else:
    logger.info(f"✅ CORS configured for origins: {allowed_origins}")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logger.info("✅ CORS middleware configured")


@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    """Rate limiting middleware"""
    # Skip rate limiting for health check and docs
    if request.url.path in ["/api/health", "/", "/api/docs", "/api/redoc", "/openapi.json"]:
        return await call_next(request)
    
    # Get client IP
    client_ip = request.client.host
    
    # Check rate limit
    if not check_rate_limit(client_ip):
        logger.warning(f"Rate limit exceeded for IP: {client_ip}")
        return JSONResponse(
            status_code=429,
            content={
                "error": "Rate limit exceeded",
                "detail": f"Maximum {RATE_LIMIT_REQUESTS} requests per {RATE_LIMIT_WINDOW} seconds"
            }
        )
    
    return await call_next(request)


@app.on_event("startup")
async def startup_event():
    """Log startup information"""
    logger.info("="*80)
    logger.info("🚀 Chamorro Chatbot API Starting Up")
    logger.info("="*80)
    logger.info(f"CORS Origins: {allowed_origins}")
    logger.info(f"Rate Limit: {RATE_LIMIT_REQUESTS} requests per {RATE_LIMIT_WINDOW} seconds")
    logger.info(f"Database: {os.getenv('DATABASE_URL', 'postgresql://localhost/chamorro_rag')}")
    if FREE_PROMO_ACTIVE:
        logger.info(f"🎄 FREE PROMO PERIOD ACTIVE until {FREE_PROMO_END_DATE}")
    logger.info("="*80)


@app.get("/", tags=["Root"])
async def root():
    """Root endpoint"""
    return {
        "message": "Chamorro Chatbot API",
        "version": "1.0.0",
        "docs": "/api/docs",
        "endpoints": {
            "chat": "POST /api/chat",
            "health": "GET /api/health"
        }
    }


@app.get("/api/health", response_model=HealthResponse, tags=["Health"])
async def health_check():
    """
    Health check endpoint
    
    Returns service status and database information.
    """
    try:
        # Try to get database stats (lazy load RAGDatabaseManager)
        # This may fail if transformers isn't installed (production optimization)
        try:
            from src.rag.manage_rag_db import RAGDatabaseManager
            manager = RAGDatabaseManager()
            chunk_count = manager._get_chunk_count()
        except ImportError as e:
            logger.warning(f"Could not load RAGDatabaseManager (transformers not installed): {e}")
            chunk_count = -1  # Indicate stats unavailable
        except Exception as e:
            logger.warning(f"Could not get chunk count: {e}")
            chunk_count = -1
        
        logger.info("Health check successful")
        return HealthResponse(
            status="healthy",
            database="connected" if chunk_count >= 0 else "stats_unavailable",
            chunks=chunk_count if chunk_count >= 0 else 0
        )
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return HealthResponse(
            status="degraded",
            database=f"error: {str(e)}",
            chunks=None
        )


@app.post("/api/chat", response_model=ChatResponse, tags=["Chat"])
async def chat(
    request: Request,
    authorization: Optional[str] = Header(None),
    # Form parameters (for multipart/form-data with files)
    message: Optional[str] = Form(None),
    mode: Optional[str] = Form(None),
    session_id: Optional[str] = Form(None),
    conversation_id: Optional[str] = Form(None),
    pending_id: Optional[str] = Form(None),  # Unique ID for cancel tracking
    file: Optional[UploadFile] = File(None)  # Renamed from 'image' to support all file types
):
    """
    Send a message to the chatbot (supports text, images, and documents)
    
    **Accepts:**
    - `application/json` for text-only messages
    - `multipart/form-data` for messages with files
    
    **Modes:**
    - `english`: General conversation (default)
    - `chamorro`: Chamorro-only immersion mode
    - `learn`: Learning mode with explanations
    
    **File Upload (Optional):**
    - **Images**: JPEG, PNG, WebP, GIF - Read and translate text in images
    - **PDF**: Extract and analyze text from PDF documents
    - **Word (.docx)**: Extract and analyze text from Word documents
    - **Text (.txt)**: Analyze plain text files
    
    **Authentication:**
    - Send `Authorization: Bearer <token>` header with Clerk JWT
    - Authentication is REQUIRED for all users
    
    **Example request (with file):**
    - Content-Type: multipart/form-data
    - message: "Translate this document"
    - mode: "english"
    - file: [file upload - PDF, Word, image, or text]
    - session_id: "session_1234567890_abc"
    """
    try:
        # Check if this is a JSON request (no file) or FormData request (with/without file)
        content_type = request.headers.get('content-type', '')
        
        if 'application/json' in content_type:
            # Parse JSON body for text-only messages
            body = await request.json()
            message = body.get('message')
            mode = body.get('mode', 'english')
            session_id = body.get('session_id')
            conversation_id = body.get('conversation_id')
            pending_id = body.get('pending_id')  # Parse pending_id from JSON
            file = None
        else:
            # FormData is already parsed by Form() parameters above
            mode = mode or 'english'
        
        # Validate required fields
        if not message:
            raise HTTPException(
                status_code=400,
                detail="Message is required"
            )
        
        # Validate mode
        valid_modes = ["english", "chamorro", "learn"]
        if mode not in valid_modes:
            logger.warning(f"Invalid mode requested: {mode}")
            raise HTTPException(
                status_code=400,
                detail=f"Invalid mode. Must be one of: {', '.join(valid_modes)}"
            )
        
        # Verify user authentication (REQUIRED)
        user_id = await verify_user(authorization)
        
        # Process file if present (images, PDFs, Word docs, text files)
        image_base64 = None
        file_url = None  # S3 URL
        document_text = None  # Extracted text from documents
        
        if file:
            logger.info(f"📁 File received: filename={file.filename}, content_type={file.content_type}")
            
            try:
                file_data = await file.read()
                logger.info(f"📦 File data read: {len(file_data)} bytes")
                
                # Process file based on type
                file_result = process_uploaded_file(
                    file_data=file_data,
                    content_type=file.content_type,
                    filename=file.filename
                )
                
                file_type = file_result['file_type']
                
                if file_type == 'image':
                    # Image: Use GPT-4o-mini Vision
                    image_base64 = file_result['image_base64']
                    logger.info(f"✅ Image processed: {len(image_base64)} chars base64")
                else:
                    # Document: Extract text and add to message
                    document_text = file_result['text_content']
                    if document_text:
                        logger.info(f"✅ Document text extracted: {len(document_text)} chars")
                    else:
                        raise ValueError("No text could be extracted from the document")
                
                # Upload to S3 (for persistence) - works for all file types
                file_url = upload_file_to_s3(
                    file_data=file_data,
                    filename=file.filename,
                    content_type=file.content_type
                )
                
                if file_url:
                    logger.info(f"✅ S3 upload successful: {file_url}")
                else:
                    logger.warning("⚠️  S3 upload failed, but continuing with extracted content")
                    
            except ValueError as e:
                # Known file processing errors
                raise HTTPException(status_code=400, detail=str(e))
            except Exception as e:
                logger.error(f"❌ Failed to process file: {e}")
                raise HTTPException(
                    status_code=400,
                    detail=f"Failed to process file: {str(e)}"
                )
        
        # If document text was extracted, append it to the message
        final_message = message
        if document_text:
            # Truncate very long documents (GPT-4o has 128k context, but let's be reasonable)
            max_doc_chars = 50000  # ~12k tokens
            if len(document_text) > max_doc_chars:
                document_text = document_text[:max_doc_chars] + "\n\n[Document truncated - showing first 50,000 characters]"
            
            final_message = f"{message}\n\n--- Document Content ---\n{document_text}"
            logger.info(f"📄 Message augmented with document text: {len(final_message)} total chars")
        
        logger.info(f"Chat request: mode={mode}, user_id={user_id or 'anonymous'}, session_id={session_id}, pending_id={pending_id}, has_file={file is not None}, has_image={image_base64 is not None}, has_doc_text={document_text is not None}")
        
        # Get response from chatbot service
        # Run in thread pool to allow cancel requests to be processed in parallel
        result = await asyncio.to_thread(
            get_chatbot_response,
            message=final_message,  # Full message with doc content for LLM
            mode=mode,
            conversation_length=0,  # For now, stateless (no session history)
            session_id=session_id,  # Pass session_id for logging
            user_id=user_id,  # Pass user_id for tracking
            conversation_id=conversation_id,  # Pass conversation_id for multi-conversation support
            image_base64=image_base64,  # Pass base64-encoded image for Vision
            image_url=file_url,  # Pass S3 URL for logging (works for all file types)
            pending_id=pending_id,  # Pass pending_id for cancel tracking
            original_message=message  # Original user message for logging/display
        )
        
        # Check if the request was cancelled - if so, return a cancelled response
        if result.get("cancelled"):
            logger.info(f"⚠️  Message {pending_id} was cancelled, returning empty response")
            return ChatResponse(
                response="Message was cancelled.",
                mode=mode,
                sources=[],
                used_rag=False,
                used_web_search=False,
                response_time=result["response_time"]
            )
        
        logger.info(f"Chatbot service called successfully with image_base64={'present' if image_base64 else 'None'}, doc_text={'present' if document_text else 'None'}")
        
        # Convert sources to SourceInfo models
        sources = [
            SourceInfo(name=s["name"], page=s["page"])
            for s in result["sources"]
        ]
        
        logger.info(f"Chat response: used_rag={result['used_rag']}, used_web_search={result['used_web_search']}, response_time={result['response_time']:.2f}s")
        
        return ChatResponse(
            response=result["response"],
            mode=mode,
            sources=sources,
            used_rag=result["used_rag"],
            used_web_search=result["used_web_search"],
            response_time=result["response_time"]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        # Log error with proper logging
        logger.error(f"Error in chat endpoint: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )


# --- Cancel Message Endpoint ---

@app.post("/api/chat/cancel/{pending_id}", tags=["Chat"])
async def cancel_chat_message(pending_id: str):
    """
    Cancel a pending chat message.
    
    This prevents the response from being saved to the database.
    Call this when the user clicks "Stop" to cancel a generating message.
    
    Args:
        pending_id: The unique ID of the pending message to cancel
        
    Returns:
        {"success": bool, "message": str}
    """
    try:
        success = cancel_pending_message(pending_id)
        
        if success:
            logger.info(f"✅ Message {pending_id} marked as cancelled")
            return {"success": True, "message": f"Message {pending_id} cancelled"}
        else:
            logger.warning(f"⚠️  Message {pending_id} was already cancelled")
            return {"success": True, "message": f"Message {pending_id} was already cancelled"}
            
    except Exception as e:
        logger.error(f"❌ Failed to cancel message {pending_id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to cancel message: {str(e)}"
        )


# --- Streaming Chat Endpoint ---

@app.post("/api/chat/stream", tags=["Chat"])
async def chat_stream(
    request: Request,
    background_tasks: BackgroundTasks,
    authorization: Optional[str] = Header(None),
    message: Optional[str] = Form(None),
    mode: Optional[str] = Form(None),
    session_id: Optional[str] = Form(None),
    conversation_id: Optional[str] = Form(None),
    pending_id: Optional[str] = Form(None),
    skill_level: Optional[str] = Form(None),  # User's skill level for personalized responses
    file: Optional[UploadFile] = File(None),  # Legacy single file support
    files: List[UploadFile] = File(default=[])  # New: multiple files support
):
    """
    Streaming chat endpoint - returns Server-Sent Events (SSE) for real-time response.
    
    **Supports up to 5 files** (images, PDFs, Word docs, text files).
    
    **Event Types:**
    - `metadata`: Sources, RAG status (sent first)
    - `chunk`: Response text chunks (sent as generated)
    - `done`: Completion signal with response_time
    - `cancelled`: User cancelled the request
    - `error`: Error occurred
    
    **Example SSE stream:**
    ```
    data: {"type": "metadata", "sources": [...], "used_rag": true}
    
    data: {"type": "chunk", "content": "Håfa "}
    
    data: {"type": "chunk", "content": "Adai "}
    
    data: {"type": "chunk", "content": "means..."}
    
    data: {"type": "done", "response_time": 2.5}
    ```
    """
    try:
        # Check if this is a JSON request or FormData request
        content_type = request.headers.get('content-type', '')
        
        if 'application/json' in content_type:
            body = await request.json()
            message = body.get('message')
            mode = body.get('mode', 'english')
            session_id = body.get('session_id')
            conversation_id = body.get('conversation_id')
            pending_id = body.get('pending_id')
            skill_level = body.get('skill_level')  # User's skill level
            files = []
            file = None
        else:
            mode = mode or 'english'
            # Combine legacy single file with new multiple files support
            if file and file.filename:
                files = [file] + list(files)
        
        # Validate required fields
        if not message:
            raise HTTPException(status_code=400, detail="Message is required")
        
        # Validate mode
        valid_modes = ["english", "chamorro", "learn"]
        if mode not in valid_modes:
            raise HTTPException(status_code=400, detail=f"Invalid mode. Must be one of: {', '.join(valid_modes)}")
        
        # Limit number of files
        MAX_FILES = 5
        if len(files) > MAX_FILES:
            raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES} files allowed")
        
        # Verify user authentication
        user_id = await verify_user(authorization)
        
        # Set Sentry context for this request
        set_user_context(user_id=user_id)
        set_request_context(
            conversation_id=conversation_id,
            mode=mode,
            model=os.getenv("CHAT_MODEL", "deepseek-v3")
        )
        
        # Process files if present
        image_base64 = None  # First image for vision model
        all_images_base64 = []  # All images for multi-image support
        document_texts = []  # Text content from documents
        final_message = message
        files_to_upload = []  # Store file data for background upload
        
        for idx, uploaded_file in enumerate(files):
            if not uploaded_file.filename:
                continue
                
            logger.info(f"📁 Stream: Processing file {idx+1}/{len(files)}: {uploaded_file.filename}")
            try:
                # Read file data once (we'll need it for both processing and S3 upload)
                file_data = await uploaded_file.read()
                
                # Store for background S3 upload
                files_to_upload.append({
                    'data': file_data,
                    'filename': uploaded_file.filename,
                    'content_type': uploaded_file.content_type,
                    'index': idx
                })
                
                # Extract content (fast - this is what the LLM needs)
                file_result = process_uploaded_file(
                    file_data=file_data,
                    content_type=uploaded_file.content_type,
                    filename=uploaded_file.filename
                )
                
                # Handle images
                if file_result.get("image_base64"):
                    all_images_base64.append(file_result["image_base64"])
                    # Use first image for the vision model
                    if image_base64 is None:
                        image_base64 = file_result["image_base64"]
                
                # Handle documents (text_content is returned by process_uploaded_file)
                if file_result.get("text_content"):
                    doc_text = file_result["text_content"]
                    document_texts.append(f"[{uploaded_file.filename}]\n{doc_text}")
                    
            except Exception as e:
                logger.error(f"File processing error for {uploaded_file.filename}: {e}")
        
        # Schedule S3 uploads as background tasks (user doesn't wait!)
        for file_info in files_to_upload:
            background_tasks.add_task(
                upload_file_to_s3_background,
                file_data=file_info['data'],
                filename=file_info['filename'],
                content_type=file_info['content_type'],
                user_id=user_id,
                conversation_id=conversation_id,
                file_index=file_info['index']
            )
        
        if files_to_upload:
            logger.info(f"📤 Scheduled {len(files_to_upload)} background S3 uploads")
        
        # Combine all document texts into the message
        if document_texts:
            max_doc_chars = 50000  # Total limit for all documents
            combined_docs = "\n\n".join(document_texts)
            if len(combined_docs) > max_doc_chars:
                combined_docs = combined_docs[:max_doc_chars] + "\n\n[Documents truncated]"
            final_message = f"{message}\n\n--- Document Content ({len(document_texts)} file{'s' if len(document_texts) > 1 else ''}) ---\n{combined_docs}"
        
        # Log file summary
        if files:
            logger.info(f"📁 Processed {len(files)} files: {len(all_images_base64)} images, {len(document_texts)} documents")
        
        # Note: image_url will be updated by background task after S3 upload completes
        file_url = None  # S3 upload happens in background
        
        logger.info(f"Stream request: mode={mode}, user_id={user_id or 'anonymous'}, pending_id={pending_id}")
        
        # Create SSE generator using queue to bridge sync/async
        async def generate_sse():
            # Use a thread-safe queue to pass events from sync thread to async generator
            event_queue: queue.Queue = queue.Queue()
            stream_done = threading.Event()
            
            def run_sync_generator():
                """Run the sync generator in a thread and put events in queue"""
                try:
                    for event in get_chatbot_response_stream(
                        message=final_message,  # Full message with doc content for LLM
                        mode=mode,
                        conversation_length=0,
                        session_id=session_id,
                        user_id=user_id,
                        conversation_id=conversation_id,
                        image_base64=image_base64,
                        image_url=file_url,
                        pending_id=pending_id,
                        original_message=message,  # Original user message for logging/display
                        skill_level=skill_level  # User's skill level for personalization
                    ):
                        event_queue.put(event)
                except Exception as e:
                    event_queue.put({"type": "error", "content": str(e)})
                finally:
                    stream_done.set()
            
            # Start the sync generator in a background thread
            thread = threading.Thread(target=run_sync_generator, daemon=True)
            thread.start()
            
            # Yield events as they come from the queue
            while True:
                # Try to get from queue without blocking
                try:
                    event = event_queue.get_nowait()
                    yield f"data: {json.dumps(event)}\n\n"
                except queue.Empty:
                    # No event available
                    if stream_done.is_set() and event_queue.empty():
                        # Stream is finished and queue is empty
                        yield "data: [DONE]\n\n"
                        break
                    # Wait a bit before checking again
                    await asyncio.sleep(0.01)
                except Exception as e:
                    logger.error(f"SSE error: {e}")
                    yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"
                    break
        
        return StreamingResponse(
            generate_sse(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"  # Disable nginx buffering
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in stream endpoint: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# --- Evaluation Endpoint (For Testing Only) ---

@app.post("/api/eval/chat", response_model=ChatResponse, tags=["Evaluation"])
async def eval_chat(
    request: Request,
    message: Optional[str] = Form(None),
    mode: Optional[str] = Form(None),
    session_id: Optional[str] = Form(None),
    conversation_id: Optional[str] = Form(None)  # NEW: For context testing
):
    """
    **EVALUATION ENDPOINT** - No authentication required.
    
    This endpoint is for automated testing and evaluation only.
    It bypasses authentication to allow automated test suites to run.
    
    **DO NOT USE THIS IN PRODUCTION CLIENT** - Use `/api/chat` instead.
    
    **Modes:**
    - `english`: General conversation (default)
    - `chamorro`: Chamorro-only immersion mode
    - `learn`: Learning mode with explanations
    
    **New: Conversation Context Testing**
    Pass `conversation_id` to test multi-turn conversations with context.
    The endpoint will retrieve and use prior messages from that conversation.
    """
    try:
        # Check if this is a JSON request
        content_type = request.headers.get('content-type', '')
        
        skill_level = None
        if 'application/json' in content_type:
            body = await request.json()
            message = body.get('message')
            mode = body.get('mode', 'english')
            session_id = body.get('session_id')
            conversation_id = body.get('conversation_id')  # NEW: For context testing
            skill_level = body.get('skill_level')  # Optional: beginner, intermediate, advanced
        else:
            mode = mode or 'english'
        
        # Validate required fields
        if not message:
            raise HTTPException(
                status_code=400,
                detail="Message is required"
            )
        
        # Validate mode
        valid_modes = ["english", "chamorro", "learn"]
        if mode not in valid_modes:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid mode. Must be one of: {', '.join(valid_modes)}"
            )
        
        # Validate skill level if provided
        valid_skill_levels = ["beginner", "intermediate", "advanced"]
        if skill_level and skill_level not in valid_skill_levels:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid skill_level. Must be one of: {', '.join(valid_skill_levels)}"
            )
        
        logger.info(f"Eval chat request: mode={mode}, skill_level={skill_level}, session_id={session_id}, conversation_id={conversation_id}")
        
        # Get response from chatbot service (no auth required for eval)
        # Now supports conversation_id for context testing!
        result = get_chatbot_response(
            message=message,
            mode=mode,
            conversation_length=0,
            session_id=session_id or f"eval_{int(time.time())}",
            user_id=None,  # No user tracking for eval
            conversation_id=conversation_id,  # NEW: Pass conversation_id for context
            image_base64=None,
            image_url=None,
            skill_level=skill_level  # Pass skill level for personalization testing
        )
        
        # Convert sources to SourceInfo models
        sources = [
            SourceInfo(name=s["name"], page=s["page"])
            for s in result["sources"]
        ]
        
        return ChatResponse(
            response=result["response"],
            mode=mode,
            sources=sources,
            used_rag=result["used_rag"],
            used_web_search=result["used_web_search"],
            response_time=result["response_time"]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in eval chat endpoint: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )


@app.post("/api/eval/conversations", response_model=ConversationResponse, tags=["Evaluation"])
async def eval_create_conversation(request: ConversationCreate):
    """
    **EVALUATION ENDPOINT** - Create conversation without authentication.
    
    Used by automated test suites to create conversations for context testing.
    
    **DO NOT USE THIS IN PRODUCTION CLIENT** - Use `/api/conversations` instead.
    """
    try:
        # Create conversation with a test user ID (so we can identify test data)
        # conversations.create_conversation returns a ConversationResponse directly
        conversation = conversations.create_conversation(
            user_id="eval_test_user",  # Special marker for test conversations
            title=request.title or "Eval Test Chat"
        )
        
        logger.info(f"Eval conversation created: {conversation.id}")
        
        # Return the ConversationResponse directly (it's already the right type)
        return conversation
        
    except Exception as e:
        logger.error(f"Error creating eval conversation: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )


# --- Conversation Management Endpoints ---

@app.post("/api/conversations", response_model=ConversationResponse, tags=["Conversations"])
async def create_conversation_endpoint(
    request: ConversationCreate,
    authorization: Optional[str] = Header(None)
):
    """
    Create a new conversation.
    
    Anonymous users can create conversations (user_id will be NULL).
    Authenticated users get conversations tied to their account.
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        # Create conversation
        conversation = conversations.create_conversation(
            user_id=user_id,
            title=request.title or "New Chat"
        )
        
        logger.info(f"Created conversation: {conversation.id} for user: {user_id or 'anonymous'}")
        return conversation
    except Exception as e:
        logger.error(f"Failed to create conversation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/conversations", response_model=ConversationListResponse, tags=["Conversations"])
async def list_conversations(
    authorization: Optional[str] = Header(None)
):
    """
    List all conversations for the authenticated user.
    
    Returns empty list for anonymous users (no auth token).
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        # Get conversations
        conversation_list = conversations.get_conversations(user_id=user_id)
        
        logger.info(f"Listed {len(conversation_list.conversations)} conversations for user: {user_id or 'anonymous'}")
        return conversation_list
    except Exception as e:
        logger.error(f"Failed to list conversations: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/init", response_model=InitResponse, tags=["Conversations"])
async def init_user_data(
    active_conversation_id: Optional[str] = None,
    authorization: Optional[str] = Header(None)
):
    """
    Initialize user data in a single request.
    
    Returns:
    - List of user's conversations
    - Messages for the active conversation (if provided and exists)
    - Active conversation ID (validated)
    
    This endpoint optimizes initial page load by combining multiple requests into one,
    eliminating the waterfall effect and reducing total latency by ~50%.
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        # Get conversations
        conversation_list = conversations.get_conversations(user_id=user_id)
        
        # Initialize response
        messages_list = []
        validated_conversation_id = None
        
        # If active_conversation_id provided, validate it exists and get messages
        if active_conversation_id:
            # Check if conversation exists in user's list
            conversation_exists = any(
                c.id == active_conversation_id 
                for c in conversation_list.conversations
            )
            
            if conversation_exists:
                # Get messages for this conversation
                messages_response = conversations.get_conversation_messages(active_conversation_id)
                messages_list = messages_response.messages
                validated_conversation_id = active_conversation_id
                logger.info(f"✅ Loaded {len(messages_list)} messages for conversation: {active_conversation_id}")
            else:
                logger.warning(f"⚠️  Conversation {active_conversation_id} not found in user's list")
        else:
            # No active_conversation_id provided = user wants a new chat
            # Return empty messages and null conversation ID
            logger.info("🆕 No active conversation provided - returning empty state for new chat")
        
        logger.info(
            f"🚀 Init complete: {len(conversation_list.conversations)} conversations, "
            f"{len(messages_list)} messages, active_id={validated_conversation_id}"
        )
        
        return InitResponse(
            conversations=conversation_list.conversations,
            messages=messages_list,
            active_conversation_id=validated_conversation_id
        )
        
    except Exception as e:
        logger.error(f"❌ Failed to initialize user data: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/conversations/{conversation_id}/messages", response_model=MessagesResponse, tags=["Conversations"])
async def get_conversation_messages_endpoint(
    conversation_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Get all messages for a specific conversation.
    
    TODO: Add ownership verification for security.
    """
    try:
        # Get messages
        messages = conversations.get_conversation_messages(conversation_id)
        
        logger.info(f"Retrieved {len(messages.messages)} messages for conversation: {conversation_id}")
        return messages
    except Exception as e:
        logger.error(f"Failed to get messages: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/api/conversations/{conversation_id}", tags=["Conversations"])
async def update_conversation_endpoint(
    conversation_id: str,
    request: ConversationCreate,
    authorization: Optional[str] = Header(None)
):
    """
    Update a conversation's title.
    
    Only the owner can update their conversation.
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        if not request.title:
            raise HTTPException(status_code=400, detail="Title is required")
        
        # Update conversation
        updated = conversations.update_conversation_title(
            conversation_id=conversation_id,
            title=request.title,
            user_id=user_id
        )
        
        if not updated:
            raise HTTPException(status_code=404, detail="Conversation not found or access denied")
        
        logger.info(f"Updated conversation: {conversation_id} title to '{request.title}' for user: {user_id or 'anonymous'}")
        return {"success": True, "message": "Conversation updated"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update conversation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/conversations/{conversation_id}", tags=["Conversations"])
async def delete_conversation_endpoint(
    conversation_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Delete a conversation (and all its messages).
    
    Only the owner can delete their conversation.
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        # Delete conversation
        deleted = conversations.delete_conversation(conversation_id, user_id=user_id)
        
        if not deleted:
            raise HTTPException(status_code=404, detail="Conversation not found or access denied")
        
        logger.info(f"Deleted conversation: {conversation_id} for user: {user_id or 'anonymous'}")
        return {"success": True, "message": "Conversation deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete conversation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/conversations/{conversation_id}/messages/after/{timestamp}", tags=["Conversations"])
async def delete_messages_after_endpoint(
    conversation_id: str,
    timestamp: int,
    authorization: Optional[str] = Header(None)
):
    """
    Delete all messages in a conversation after a given timestamp.
    
    Used for the Edit & Regenerate feature. When a user edits a previous message,
    all subsequent messages are deleted and a new response is generated.
    
    Args:
        conversation_id: The conversation ID
        timestamp: Unix timestamp in milliseconds - delete messages after this time
    """
    try:
        # Verify authentication
        user_id = await verify_user(authorization)
        
        # Delete messages after the timestamp
        deleted_count = conversations.delete_messages_after(
            conversation_id, 
            timestamp, 
            user_id=user_id
        )
        
        logger.info(f"Deleted {deleted_count} messages after {timestamp} in conversation {conversation_id} (Edit & Regenerate)")
        return {"success": True, "deleted_count": deleted_count}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete messages after timestamp: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/conversations/system-message", tags=["Conversations"])
async def create_system_message_endpoint(
    request: SystemMessageCreate,
    authorization: Optional[str] = Header(None)
):
    """
    Create a system message (e.g., mode change indicator).
    
    System messages are used to track events like mode switching for analytics
    and to provide context in conversation history.
    """
    try:
        # Verify authentication (optional - allow anonymous users too)
        user_id = None
        try:
            user_id = await verify_user(authorization)
        except:
            pass  # Allow anonymous users
        
        # Get session ID from somewhere (we'll need to add this to the request model)
        session_id = None
        # TODO: Consider adding session_id to SystemMessageCreate if needed
        
        # Create system message
        success = conversations.create_system_message(
            conversation_id=request.conversation_id,
            content=request.content,
            mode=request.mode,
            user_id=user_id,
            session_id=session_id
        )
        
        if not success:
            raise HTTPException(status_code=500, detail="Failed to create system message")
        
        logger.info(f"Created system message in conversation {request.conversation_id}: {request.content}")
        return {"success": True, "message": "System message created"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create system message: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Share Conversation Endpoints ---

@app.post("/api/conversations/{conversation_id}/share", response_model=ShareConversationResponse, tags=["Share"])
async def create_share_link(
    conversation_id: str,
    request: ShareConversationRequest = None,
    authorization: Optional[str] = Header(None)
):
    """
    Create a shareable public link for a conversation.
    
    The owner can share their conversation with anyone via a unique URL.
    Optionally set an expiration date.
    """
    try:
        # Verify user owns this conversation
        user_id = await verify_user(authorization)
        
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Verify conversation exists and belongs to user
        cursor.execute("""
            SELECT id, user_id FROM conversations 
            WHERE id = %s AND deleted_at IS NULL
        """, (conversation_id,))
        conv = cursor.fetchone()
        
        if not conv:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Conversation not found")
        
        if conv[1] != user_id:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=403, detail="You don't own this conversation")
        
        # Check if share already exists for this conversation
        cursor.execute("""
            SELECT share_id, expires_at, created_at FROM shared_conversations
            WHERE conversation_id = %s AND user_id = %s
        """, (conversation_id, user_id))
        existing = cursor.fetchone()
        
        if existing:
            # Return existing share
            share_id = existing[0]
            expires_at = existing[1]
            created_at = existing[2]
            cursor.close()
            conn.close()
            
            # Build share URL (frontend URL)
            frontend_url = os.getenv("FRONTEND_URL", "https://hafagpt.com")
            share_url = f"{frontend_url}/share/{share_id}"
            
            return ShareConversationResponse(
                share_id=share_id,
                share_url=share_url,
                expires_at=expires_at,
                created_at=created_at
            )
        
        # Create new share
        import uuid
        share_id = str(uuid.uuid4())[:8]  # Short ID for nicer URLs
        internal_id = str(uuid.uuid4())
        
        # Calculate expiration if specified
        expires_at = None
        if request and request.expires_in_days:
            from datetime import timedelta
            expires_at = datetime.now() + timedelta(days=request.expires_in_days)
        
        cursor.execute("""
            INSERT INTO shared_conversations (id, share_id, conversation_id, user_id, expires_at)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING created_at
        """, (internal_id, share_id, conversation_id, user_id, expires_at))
        
        created_at = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()
        
        # Build share URL
        frontend_url = os.getenv("FRONTEND_URL", "https://hafagpt.com")
        share_url = f"{frontend_url}/share/{share_id}"
        
        logger.info(f"🔗 Created share link for conversation {conversation_id}: {share_url}")
        
        return ShareConversationResponse(
            share_id=share_id,
            share_url=share_url,
            expires_at=expires_at,
            created_at=created_at
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create share link: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/share/{share_id}", response_model=SharedConversationResponse, tags=["Share"])
async def get_shared_conversation(share_id: str):
    """
    Get a shared conversation by its public share ID.
    
    This is a PUBLIC endpoint - no authentication required.
    Anyone with the link can view the conversation.
    """
    try:
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get share info and conversation
        cursor.execute("""
            SELECT 
                sc.share_id,
                sc.conversation_id,
                sc.expires_at,
                sc.view_count,
                c.title,
                c.created_at
            FROM shared_conversations sc
            JOIN conversations c ON c.id = sc.conversation_id
            WHERE sc.share_id = %s
        """, (share_id,))
        
        share = cursor.fetchone()
        
        if not share:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Share not found")
        
        # Check expiration
        expires_at = share[2]
        if expires_at and expires_at < datetime.now(expires_at.tzinfo):
            cursor.close()
            conn.close()
            raise HTTPException(status_code=410, detail="This share link has expired")
        
        conversation_id = share[1]
        view_count = share[3]
        title = share[4]
        created_at = share[5]
        
        # Increment view count
        cursor.execute("""
            UPDATE shared_conversations 
            SET view_count = view_count + 1 
            WHERE share_id = %s
        """, (share_id,))
        conn.commit()
        
        cursor.close()
        conn.close()
        
        # Get messages using existing function
        messages_response = conversations.get_conversation_messages(conversation_id)
        
        return SharedConversationResponse(
            share_id=share_id,
            title=title,
            created_at=created_at,
            messages=messages_response.messages,
            view_count=view_count + 1  # Include the current view
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get shared conversation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/conversations/{conversation_id}/share", tags=["Share"])
async def revoke_share_link(
    conversation_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Revoke (delete) a share link for a conversation.
    
    Only the owner can revoke their share.
    """
    try:
        user_id = await verify_user(authorization)
        
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Delete share if user owns it
        cursor.execute("""
            DELETE FROM shared_conversations
            WHERE conversation_id = %s AND user_id = %s
            RETURNING share_id
        """, (conversation_id, user_id))
        
        deleted = cursor.fetchone()
        conn.commit()
        cursor.close()
        conn.close()
        
        if not deleted:
            raise HTTPException(status_code=404, detail="Share not found or you don't own it")
        
        logger.info(f"🗑️ Revoked share link {deleted[0]} for conversation {conversation_id}")
        
        return {"success": True, "message": "Share link revoked"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to revoke share link: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/conversations/{conversation_id}/share", response_model=ShareInfoResponse, tags=["Share"])
async def get_share_info(
    conversation_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Get share info for a conversation (if shared).
    
    Only the owner can see their share info.
    """
    try:
        user_id = await verify_user(authorization)
        
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT 
                sc.id,
                sc.share_id,
                sc.conversation_id,
                sc.created_at,
                sc.expires_at,
                sc.view_count,
                c.title
            FROM shared_conversations sc
            JOIN conversations c ON c.id = sc.conversation_id
            WHERE sc.conversation_id = %s AND sc.user_id = %s
        """, (conversation_id, user_id))
        
        share = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not share:
            raise HTTPException(status_code=404, detail="No share exists for this conversation")
        
        frontend_url = os.getenv("FRONTEND_URL", "https://hafagpt.com")
        share_url = f"{frontend_url}/share/{share[1]}"
        
        return ShareInfoResponse(
            id=share[0],
            share_id=share[1],
            share_url=share_url,
            conversation_id=share[2],
            conversation_title=share[6],
            created_at=share[3],
            expires_at=share[4],
            view_count=share[5]
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get share info: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Text-to-Speech Endpoint ---

@app.post("/api/tts", tags=["Speech"])
async def text_to_speech(
    text: str = Form(...),
    voice: str = Form(default="shimmer")  # Shimmer works best for Chamorro/Spanish
):
    """
    Convert text to speech using OpenAI TTS HD.
    
    Returns base64-encoded MP3 audio that can be played in browser.
    
    Using tts-1-hd model for higher quality pronunciation.
    
    Voices available (try different ones for Chamorro):
    - shimmer: Soft, gentle (BEST for Spanish/Chamorro) ⭐
    - alloy: Neutral, balanced (good for multilingual)
    - echo: Clear, professional  
    - fable: Expressive, dramatic
    - onyx: Deep, authoritative (male voice)
    - nova: Warm, engaging (female voice)
    """
    try:
        # Import OpenAI client (lazy load to avoid startup overhead)
        from openai import OpenAI
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        client = OpenAI(api_key=api_key)
        
        # Limit text length (OpenAI TTS max is 4096 characters)
        text_to_speak = text[:4096]
        
        # Note: OpenAI TTS handles Chamorro reasonably well without hints
        # The model auto-detects language from text content
        
        logger.info(f"🔊 TTS request: {len(text_to_speak)} characters, voice={voice}")
        
        # Call OpenAI TTS API
        response = client.audio.speech.create(
            model="tts-1",  # Standard quality (2x faster, $0.015/1K chars) - HD not needed for single words
            voice=voice,
            input=text_to_speak,
        )
        
        # Get audio bytes
        audio_bytes = response.content
        
        # Convert to base64 for easy transport
        audio_base64 = base64.b64encode(audio_bytes).decode('utf-8')
        
        logger.info(f"✅ TTS successful: {len(audio_bytes)} bytes")
        
        return {
            "audio": audio_base64,
            "format": "mp3"
        }
        
    except ImportError:
        logger.error("❌ OpenAI library not installed")
        raise HTTPException(status_code=500, detail="OpenAI library not available")
    except Exception as e:
        logger.error(f"❌ TTS failed: {e}")
        raise HTTPException(status_code=500, detail=f"TTS generation failed: {str(e)}")


# --- Flashcard Generation Endpoint ---

@app.post("/api/generate-flashcards", response_model=GenerateFlashcardsResponse, tags=["Flashcards"])
async def generate_flashcards(
    topic: str = Form(...),
    count: int = Form(default=20),
    variety: str = Form(default="basic"),  # New parameter: basic, conversational, or advanced
    previous_cards: str = Form(default="[]"),  # JSON string of previously generated cards
):
    """
    Generate Chamorro language flashcards using GPT + RAG knowledge.
    
    No database storage - generates fresh cards each time (stateless MVP).
    
    Available topics:
    - greetings: Greetings & basic phrases
    - family: Family members & relationships
    - food: Food & cooking vocabulary
    - numbers: Numbers 1-20
    - verbs: Common action verbs
    - common-phrases: Everyday phrases
    
    Args:
        topic: Topic category for flashcards
        count: Number of cards to generate (default: 20, max: 30)
        variety: Card variety level (basic, conversational, advanced)
    
    Returns:
        JSON with array of flashcards
    """
    import time
    start_time = time.time()
    logger.info(f"🎴 [FLASHCARDS] Request received - topic: {topic}, count: {count}, variety: {variety}")
    
    try:
        # Parse previous cards if provided
        previous_cards_list = []
        try:
            previous_cards_list = json.loads(previous_cards)
            if previous_cards_list:
                logger.info(f"🎴 [FLASHCARDS] Received {len(previous_cards_list)} previous cards to avoid duplicates")
        except json.JSONDecodeError:
            logger.warning(f"🎴 [FLASHCARDS] Failed to parse previous_cards, ignoring")
        
        # Validate topic
        valid_topics = ["greetings", "family", "food", "numbers", "verbs", "common-phrases"]
        if topic not in valid_topics:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid topic. Must be one of: {', '.join(valid_topics)}"
            )
        
        # Limit count
        count = min(count, 30)
        
        logger.info(f"🎴 [FLASHCARDS] Starting RAG search...")
        rag_start = time.time()
        
        # Import RAG and OpenAI
        from src.rag.chamorro_rag import rag
        from openai import OpenAI
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        client = OpenAI(api_key=api_key)
        
        # Use RAG to get relevant content for the topic
        topic_queries = {
            "greetings": "Chamorro greetings, hello, good morning, good night, how are you",
            "family": "Chamorro family members, mother, father, brother, sister, relatives",
            "food": "Chamorro food vocabulary, cooking, eating, meals, ingredients",
            "numbers": "Chamorro numbers, counting, one through twenty",
            "verbs": "Chamorro action verbs, common verbs, doing words",
            "common-phrases": "Chamorro everyday phrases, useful expressions, conversations"
        }
        
        # Map topics to card types for source prioritization
        topic_to_card_type = {
            "greetings": "phrases",         # Greetings are conversational
            "family": "words",               # Family members are vocabulary
            "food": "words",                 # Food items are vocabulary
            "numbers": "numbers",            # Numbers get special treatment
            "verbs": "words",                # Verbs are vocabulary
            "common-phrases": "phrases"      # Phrases need conversational context
        }
        
        query = topic_queries.get(topic, f"Chamorro {topic} vocabulary")
        card_type = topic_to_card_type.get(topic, "words")  # Default to words
        
        logger.info(f"🎴 [FLASHCARDS] Card type: {card_type} (will prioritize appropriate sources)")
        
        # Search RAG database for relevant content with card-type specific prioritization
        rag_context, rag_sources = rag.create_context(query, k=5, card_type=card_type)  # Pass card_type!
        
        rag_end = time.time()
        logger.info(f"🎴 [FLASHCARDS] RAG search took: {(rag_end - rag_start):.2f}s")
        logger.info(f"🎴 [FLASHCARDS] Retrieved {len(rag_sources)} sources")
        
        # Create GPT prompt with variety-specific instructions
        logger.info(f"🎴 [FLASHCARDS] Building GPT prompt with variety: {variety}...")
        
        # Variety-specific instructions
        variety_instructions = {
            "basic": """- Focus on BASIC everyday usage and common expressions
- Use simple, direct translations
- Prefer the most frequently used words/phrases
- Keep examples straightforward and practical""",
            "conversational": """- Focus on CONVERSATIONAL variations and contextual usage
- Include phrases you'd hear in daily conversations
- Show how words are used in real-life situations
- Add some colloquial or informal expressions""",
            "advanced": """- Focus on ADVANCED expressions and nuanced meanings
- Include less common variations or regional differences
- Show more complex sentence structures
- Include idiomatic or cultural expressions"""
        }
        
        variety_instruction = variety_instructions.get(variety, variety_instructions["basic"])
        
        # Build previous cards section if any exist
        previous_cards_section = ""
        if previous_cards_list:
            previous_cards_text = "\n".join([
                f"  • {card.get('front', '')} = {card.get('back', '')}" 
                for card in previous_cards_list
            ])
            previous_cards_section = f"""
⚠️ CRITICAL - AVOID ALL DUPLICATES ⚠️

You have ALREADY generated these {len(previous_cards_list)} cards:
{previous_cards_text}

DO NOT create cards that:
1. Use the SAME Chamorro word/phrase (case-insensitive)
2. Use the SAME English translation (case-insensitive)
3. Have SIMILAR or OVERLAPPING meanings
4. Are SYNONYMS or NEAR-SYNONYMS of previous cards

MUST GENERATE: {count} COMPLETELY NEW, UNIQUE, DIFFERENT cards.
- Choose DIFFERENT Chamorro words/phrases NOT in the list above
- Choose DIFFERENT English meanings NOT in the list above
- Explore DIFFERENT aspects of the topic
"""
        
        prompt = f"""You are a Chamorro language teacher creating educational flashcards.

Generate {count} HIGH-QUALITY, UNIQUE Chamorro language flashcards for the topic: {topic}

IMPORTANT: Generate DIFFERENT content than what's typically found in basic dictionaries or common phrasebooks.
Avoid the most obvious/common phrases like "Håfa Adai", "Si Yu'os Ma'åse'", "Maolek ha' yu'" unless specifically relevant.

Use the reference materials provided below to create accurate flashcards with proper translations and pronunciation.

REFERENCE MATERIALS:
{rag_context}

VARIETY LEVEL: {variety.upper()}
{variety_instruction}

{previous_cards_section}

REQUIREMENTS:
1. Each flashcard should have:
   - front: Chamorro word or phrase
   - back: Clear English translation
   - pronunciation: Phonetic guide (e.g., "HAH-fah ah-DYE")
   - example: A simple example sentence showing usage in Chamorro
   - category: "{topic}"

2. Focus on UNIQUE and INTERESTING content (not just the most basic phrases)
3. Use proper Chamorro spelling (å, ñ characters)
4. Provide accurate pronunciation guides
5. Keep examples simple but authentic
6. Ensure each card is DISTINCT and offers new learning value

Return ONLY a valid JSON array with this exact structure:
[
  {{
    "front": "Kao un gof guaiya este?",
    "back": "Do you really like this?",
    "pronunciation": "kah-oh oon goff gwah-EE-yah EH-steh",
    "example": "Kao un gof guaiya este kandet? (Do you really like this candy?)",
    "category": "{topic}"
  }}
]

Generate exactly {count} flashcards. Return only the JSON array, no other text."""

        # Call GPT-4o-mini
        logger.info("🎴 [FLASHCARDS] Calling GPT-4o-mini...")
        gpt_start = time.time()
        
        response = client.chat.completions.create(
            model="gpt-4o",  # Premium model: 96% accuracy, perfect grammar, faster than 4o-mini
            messages=[
                {
                    "role": "system",
                    "content": "You are a Chamorro language expert creating educational flashcards. Always respond with valid JSON arrays only."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,  # Some creativity for variety
            max_tokens=1500  # Reduced from 3000 - sufficient for 3 cards
        )
        
        gpt_end = time.time()
        logger.info(f"🎴 [FLASHCARDS] GPT call took: {(gpt_end - gpt_start):.2f}s")
        
        # Parse GPT response
        logger.info(f"🎴 [FLASHCARDS] Parsing GPT response...")
        parse_start = time.time()
        
        response_text = response.choices[0].message.content.strip()
        
        # Clean up response (remove markdown code blocks if present)
        if response_text.startswith("```"):
            # Remove ```json and ``` markers
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        
        response_text = response_text.strip()
        
        # Parse JSON
        flashcards_data = json.loads(response_text)
        
        # Deduplicate flashcards (safety net)
        # Pre-populate with previous cards to check against ALL cards (not just current batch)
        seen_fronts = set()
        seen_backs = set()
        
        # Add previous cards to seen sets
        for prev_card in previous_cards_list:
            prev_front = prev_card.get('front', '').lower().strip()
            prev_back = prev_card.get('back', '').lower().strip()
            if prev_front:
                seen_fronts.add(prev_front)
            if prev_back:
                seen_backs.add(prev_back)
        
        logger.info(f"🎴 [FLASHCARDS] Checking for duplicates against {len(seen_fronts)} previous fronts and {len(seen_backs)} previous backs")
        
        unique_flashcards_data = []
        duplicates_found = 0
        
        for card_data in flashcards_data:
            front_lower = card_data.get('front', '').lower().strip()
            back_lower = card_data.get('back', '').lower().strip()
            
            # Check if we've seen this card before (in previous batches OR current batch)
            if front_lower and back_lower and front_lower not in seen_fronts and back_lower not in seen_backs:
                seen_fronts.add(front_lower)
                seen_backs.add(back_lower)
                unique_flashcards_data.append(card_data)
            else:
                duplicates_found += 1
                logger.warning(f"🎴 [FLASHCARDS] ❌ Skipping duplicate card: '{card_data.get('front', '')}' = '{card_data.get('back', '')}'")
        
        logger.info(f"🎴 [FLASHCARDS] ✅ Kept {len(unique_flashcards_data)} unique cards out of {len(flashcards_data)} generated ({duplicates_found} duplicates removed)")
        
        # Validate and convert to FlashcardResponse objects
        flashcards = []
        for card_data in unique_flashcards_data:
            flashcard = FlashcardResponse(
                front=card_data.get("front", ""),
                back=card_data.get("back", ""),
                pronunciation=card_data.get("pronunciation"),
                example=card_data.get("example"),
                category=card_data.get("category", topic)
            )
            flashcards.append(flashcard)
        
        parse_end = time.time()
        logger.info(f"🎴 [FLASHCARDS] Parsing took: {(parse_end - parse_start):.2f}s")
        
        total_time = time.time() - start_time
        logger.info(f"🎴 [FLASHCARDS] ✅ Total request time: {total_time:.2f}s")
        logger.info(f"🎴 [FLASHCARDS] Breakdown: RAG={rag_end-rag_start:.1f}s, GPT={gpt_end-gpt_start:.1f}s, Parse={parse_end-parse_start:.1f}s")
        
        return GenerateFlashcardsResponse(
            flashcards=flashcards,
            topic=topic,
            count=len(flashcards)
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"❌ Failed to parse GPT response as JSON: {e}")
        logger.error(f"Response was: {response_text[:500]}")
        raise HTTPException(
            status_code=500,
            detail="Failed to generate flashcards: Invalid response format"
        )
    except Exception as e:
        logger.error(f"❌ Flashcard generation failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Flashcard generation failed: {str(e)}"
        )


# ===========================
# Flashcard Progress Endpoints
# ===========================

@app.post("/api/flashcards/decks", response_model=SaveDeckResponse, tags=["Flashcard Progress"])
async def save_flashcard_deck(request: SaveDeckRequest):
    """
    Save a deck of flashcards to the user's collection.
    
    This allows users to save custom AI-generated cards or create their own decks.
    
    Args:
        request: SaveDeckRequest with user_id, topic, title, card_type, and cards
    
    Returns:
        SaveDeckResponse with the created deck_id
    """
    import psycopg
    from datetime import datetime
    import uuid
    
    logger.info(f"💾 [SAVE DECK] User {request.user_id} saving deck: {request.title} ({len(request.cards)} cards)")
    
    try:
        conn = psycopg.connect(os.getenv("DATABASE_URL"))
        cursor = conn.cursor()
        
        # Create the deck
        deck_id = str(uuid.uuid4())
        cursor.execute(
            """
            INSERT INTO flashcard_decks (id, user_id, topic, title, card_type, created_at)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id
            """,
            (deck_id, request.user_id, request.topic, request.title, request.card_type, datetime.now())
        )
        
        # Insert all cards
        for card in request.cards:
            cursor.execute(
                """
                INSERT INTO flashcards (deck_id, front, back, pronunciation, example, created_at)
                VALUES (%s, %s, %s, %s, %s, %s)
                """,
                (deck_id, card.front, card.back, card.pronunciation, card.example, datetime.now())
            )
        
        conn.commit()
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [SAVE DECK] Successfully saved deck {deck_id} with {len(request.cards)} cards")
        
        return SaveDeckResponse(
            deck_id=deck_id,
            message=f"Successfully saved {len(request.cards)} cards to '{request.title}'"
        )
        
    except Exception as e:
        logger.error(f"❌ [SAVE DECK] Failed to save deck: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save deck: {str(e)}"
        )


@app.get("/api/flashcards/decks", response_model=UserDecksResponse, tags=["Flashcard Progress"])
async def get_user_decks(user_id: str):
    """
    Get all flashcard decks for a user with progress statistics.
    
    Returns each deck with:
    - Total cards
    - Cards reviewed (at least once)
    - Cards due for review today
    
    Args:
        user_id: User ID from Clerk
    
    Returns:
        UserDecksResponse with list of decks
    """
    import psycopg
    from datetime import datetime
    
    logger.info(f"📚 [GET DECKS] Fetching decks for user: {user_id}")
    
    try:
        conn = psycopg.connect(os.getenv("DATABASE_URL"))
        cursor = conn.cursor()
        
        # Get all decks for the user with stats
        cursor.execute(
            """
            SELECT 
                d.id,
                d.topic,
                d.title,
                d.card_type,
                d.created_at,
                COUNT(f.id) as total_cards,
                COUNT(p.flashcard_id) as cards_reviewed,
                COUNT(CASE WHEN p.next_review <= %s THEN 1 END) as cards_due
            FROM flashcard_decks d
            LEFT JOIN flashcards f ON d.id = f.deck_id
            LEFT JOIN user_flashcard_progress p ON f.id = p.flashcard_id AND p.user_id = %s
            WHERE d.user_id = %s
            GROUP BY d.id, d.topic, d.title, d.card_type, d.created_at
            ORDER BY d.created_at DESC
            """,
            (datetime.now(), user_id, user_id)
        )
        
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        
        decks = []
        for row in rows:
            decks.append(UserDeckResponse(
                id=str(row[0]),
                topic=row[1],
                title=row[2],
                card_type=row[3],
                created_at=row[4],
                total_cards=row[5] or 0,
                cards_reviewed=row[6] or 0,
                cards_due=row[7] or 0
            ))
        
        logger.info(f"✅ [GET DECKS] Found {len(decks)} decks for user {user_id}")
        
        return UserDecksResponse(decks=decks)
        
    except Exception as e:
        logger.error(f"❌ [GET DECKS] Failed to fetch decks: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to fetch decks: {str(e)}"
        )


@app.get("/api/flashcards/decks/{deck_id}/cards", response_model=DeckCardsResponse, tags=["Flashcard Progress"])
async def get_deck_cards(deck_id: str, user_id: str):
    """
    Get all cards in a deck with user progress.
    
    Args:
        deck_id: UUID of the deck
        user_id: User ID from Clerk
    
    Returns:
        DeckCardsResponse with cards and progress
    """
    import psycopg
    
    logger.info(f"🃏 [GET CARDS] Fetching cards for deck: {deck_id}, user: {user_id}")
    
    try:
        conn = psycopg.connect(os.getenv("DATABASE_URL"))
        cursor = conn.cursor()
        
        # Get deck info
        cursor.execute(
            "SELECT topic, title FROM flashcard_decks WHERE id = %s AND user_id = %s",
            (deck_id, user_id)
        )
        deck_row = cursor.fetchone()
        
        if not deck_row:
            raise HTTPException(status_code=404, detail="Deck not found")
        
        topic, title = deck_row
        
        # Get cards with progress
        cursor.execute(
            """
            SELECT 
                f.id,
                f.front,
                f.back,
                f.pronunciation,
                f.example,
                p.times_reviewed,
                p.last_reviewed,
                p.next_review,
                p.confidence
            FROM flashcards f
            LEFT JOIN user_flashcard_progress p ON f.id = p.flashcard_id AND p.user_id = %s
            WHERE f.deck_id = %s
            ORDER BY f.created_at
            """,
            (user_id, deck_id)
        )
        
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        
        cards = []
        for row in rows:
            progress = None
            if row[5] is not None:  # Has progress
                progress = FlashcardProgressInfo(
                    times_reviewed=row[5],
                    last_reviewed=row[6],
                    next_review=row[7],
                    confidence=row[8]
                )
            
            cards.append(FlashcardWithProgress(
                id=str(row[0]),
                front=row[1],
                back=row[2],
                pronunciation=row[3],
                example=row[4],
                progress=progress
            ))
        
        logger.info(f"✅ [GET CARDS] Found {len(cards)} cards in deck {deck_id}")
        
        return DeckCardsResponse(
            deck_id=deck_id,
            topic=topic,
            title=title,
            cards=cards
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [GET CARDS] Failed to fetch cards: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to fetch cards: {str(e)}"
        )


@app.post("/api/flashcards/review", response_model=ReviewCardResponse, tags=["Flashcard Progress"])
async def review_flashcard(request: ReviewCardRequest):
    """
    Mark a flashcard as reviewed and update spaced repetition schedule.
    
    Confidence levels:
    - 1 (Hard): Review tomorrow (1 day)
    - 2 (Good): Review in 1 week (7 days)
    - 3 (Easy): Review in 1 month (30 days)
    
    Args:
        request: ReviewCardRequest with user_id, flashcard_id, and confidence
    
    Returns:
        ReviewCardResponse with next review date
    """
    import psycopg
    from datetime import datetime, timedelta
    
    logger.info(f"✍️ [REVIEW] User {request.user_id} reviewed card {request.flashcard_id} with confidence {request.confidence}")
    
    try:
        # Calculate next review date based on confidence
        intervals = {
            1: 1,   # Hard: 1 day
            2: 7,   # Good: 7 days
            3: 30   # Easy: 30 days
        }
        days = intervals[request.confidence]
        next_review = datetime.now() + timedelta(days=days)
        
        conn = psycopg.connect(os.getenv("DATABASE_URL"))
        cursor = conn.cursor()
        
        # Upsert progress record
        cursor.execute(
            """
            INSERT INTO user_flashcard_progress 
                (user_id, flashcard_id, times_reviewed, last_reviewed, next_review, confidence, created_at, updated_at)
            VALUES 
                (%s, %s, 1, %s, %s, %s, %s, %s)
            ON CONFLICT (user_id, flashcard_id)
            DO UPDATE SET
                times_reviewed = user_flashcard_progress.times_reviewed + 1,
                last_reviewed = %s,
                next_review = %s,
                confidence = %s,
                updated_at = %s
            """,
            (
                request.user_id, request.flashcard_id, datetime.now(), next_review, request.confidence,
                datetime.now(), datetime.now(),
                datetime.now(), next_review, request.confidence, datetime.now()
            )
        )
        
        conn.commit()
        cursor.close()
        conn.close()
        
        # Create feedback message
        messages = {
            1: f"Let's practice this again tomorrow! 💪",
            2: f"Good job! See you in a week! 👍",
            3: f"Awesome! You've mastered this! See you in a month! 🎉"
        }
        
        logger.info(f"✅ [REVIEW] Updated progress for card {request.flashcard_id}, next review in {days} days")
        
        return ReviewCardResponse(
            next_review=next_review,
            message=messages[request.confidence],
            days_until_next=days
        )
        
    except Exception as e:
        logger.error(f"❌ [REVIEW] Failed to update progress: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to update progress: {str(e)}"
        )


# ==========================================
# Message Feedback Endpoint
# ==========================================

@app.post("/api/feedback", response_model=MessageFeedbackResponse)
async def submit_feedback(request: MessageFeedbackRequest, authorization: Optional[str] = Header(None)):
    """
    Submit feedback (thumbs up/down) on a message.
    
    Allows both authenticated and anonymous users to provide feedback.
    """
    try:
        # Try to verify user, but allow anonymous feedback
        user_id = None
        try:
            user_id = await verify_user(authorization)
        except HTTPException:
            # Anonymous user - that's okay for feedback
            logger.info("📊 [FEEDBACK] Anonymous user submitting feedback")
        
        # Validate feedback type
        if request.feedback_type not in ['up', 'down']:
            raise HTTPException(
                status_code=400,
                detail="feedback_type must be 'up' or 'down'"
            )
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Insert feedback
        cursor.execute("""
            INSERT INTO message_feedback (
                message_id, conversation_id, user_id, feedback_type, user_query, bot_response
            )
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (
            request.message_id,
            request.conversation_id,
            user_id,
            request.feedback_type,
            request.user_query,
            request.bot_response
        ))
        
        feedback_id = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [FEEDBACK] {request.feedback_type} - user: {user_id or 'anonymous'}, message: {request.message_id}")
        
        return MessageFeedbackResponse(
            status="success",
            feedback_id=str(feedback_id)
        )
        
    except Exception as e:
        logger.error(f"❌ [FEEDBACK] Failed to save feedback: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save feedback: {str(e)}"
        )


# ==========================================
# Quiz Results Endpoints
# ==========================================

@app.post("/api/quiz/results", response_model=QuizResultResponse, tags=["Quiz"])
async def save_quiz_result(
    request: QuizResultCreate,
    authorization: Optional[str] = Header(None)
):
    """
    Save a quiz result for the authenticated user.
    
    Requires authentication. Optionally includes individual question answers.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Calculate percentage
        percentage = (request.score / request.total) * 100 if request.total > 0 else 0
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Insert quiz result
        cursor.execute("""
            INSERT INTO quiz_results (
                user_id, category_id, category_title, score, total, percentage, time_spent_seconds
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id, created_at
        """, (
            user_id,
            request.category_id,
            request.category_title,
            request.score,
            request.total,
            percentage,
            request.time_spent_seconds
        ))
        
        result = cursor.fetchone()
        result_id = result[0]
        created_at = result[1]
        
        # Insert individual answers if provided
        if request.answers:
            for answer in request.answers:
                cursor.execute("""
                    INSERT INTO quiz_answers (
                        quiz_result_id, question_id, question_text, question_type,
                        user_answer, correct_answer, is_correct, explanation
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    result_id,
                    answer.question_id,
                    answer.question_text,
                    answer.question_type,
                    answer.user_answer,
                    answer.correct_answer,
                    answer.is_correct,
                    answer.explanation
                ))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        answers_count = len(request.answers) if request.answers else 0
        logger.info(f"✅ [QUIZ] Saved result for user {user_id}: {request.score}/{request.total} ({percentage:.1f}%) in {request.category_id} with {answers_count} answers")
        
        return QuizResultResponse(
            id=str(result_id),
            category_id=request.category_id,
            category_title=request.category_title,
            score=request.score,
            total=request.total,
            percentage=percentage,
            time_spent_seconds=request.time_spent_seconds,
            created_at=created_at
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [QUIZ] Failed to save result: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save quiz result: {str(e)}"
        )


@app.get("/api/quiz/results/{result_id}", response_model=QuizResultDetailResponse, tags=["Quiz"])
async def get_quiz_result_detail(
    result_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Get detailed quiz result with individual question answers.
    
    Requires authentication. User can only view their own quiz results.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get quiz result (verify ownership)
        cursor.execute("""
            SELECT id, category_id, category_title, score, total, percentage, time_spent_seconds, created_at
            FROM quiz_results
            WHERE id = %s AND user_id = %s
        """, (result_id, user_id))
        
        result_row = cursor.fetchone()
        if not result_row:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Quiz result not found")
        
        # Get individual answers
        cursor.execute("""
            SELECT id, question_id, question_text, question_type, user_answer, correct_answer, is_correct, explanation
            FROM quiz_answers
            WHERE quiz_result_id = %s
            ORDER BY created_at ASC
        """, (result_id,))
        
        answer_rows = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Build response
        answers = [
            QuizAnswerResponse(
                id=str(row[0]),
                question_id=row[1],
                question_text=row[2],
                question_type=row[3],
                user_answer=row[4],
                correct_answer=row[5],
                is_correct=row[6],
                explanation=row[7]
            )
            for row in answer_rows
        ]
        
        return QuizResultDetailResponse(
            id=str(result_row[0]),
            category_id=result_row[1],
            category_title=result_row[2],
            score=result_row[3],
            total=result_row[4],
            percentage=float(result_row[5]),
            time_spent_seconds=result_row[6],
            created_at=result_row[7],
            answers=answers
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [QUIZ] Failed to get result detail: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get quiz result: {str(e)}"
        )


@app.get("/api/quiz/stats", response_model=QuizStatsResponse, tags=["Quiz"])
async def get_quiz_stats(authorization: Optional[str] = Header(None)):
    """
    Get quiz statistics for the authenticated user.
    
    Returns total quizzes, average score, best category, and recent results.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get total quizzes and average score
        cursor.execute("""
            SELECT 
                COUNT(*) as total_quizzes,
                COALESCE(AVG(percentage), 0) as average_score
            FROM quiz_results
            WHERE user_id = %s
        """, (user_id,))
        
        stats = cursor.fetchone()
        total_quizzes = stats[0]
        average_score = float(stats[1])
        
        # Get best category (category with highest average score, min 2 attempts)
        cursor.execute("""
            SELECT 
                category_id,
                category_title,
                AVG(percentage) as avg_percentage,
                COUNT(*) as attempts
            FROM quiz_results
            WHERE user_id = %s
            GROUP BY category_id, category_title
            HAVING COUNT(*) >= 1
            ORDER BY avg_percentage DESC
            LIMIT 1
        """, (user_id,))
        
        best_category_row = cursor.fetchone()
        best_category = None
        best_category_title = None
        best_category_percentage = None
        
        if best_category_row:
            best_category = best_category_row[0]
            best_category_title = best_category_row[1]
            best_category_percentage = float(best_category_row[2])
        
        # Get recent results (last 10)
        cursor.execute("""
            SELECT id, category_id, category_title, score, total, percentage, time_spent_seconds, created_at
            FROM quiz_results
            WHERE user_id = %s
            ORDER BY created_at DESC
            LIMIT 10
        """, (user_id,))
        
        recent_rows = cursor.fetchall()
        recent_results = [
            QuizResultResponse(
                id=str(row[0]),
                category_id=row[1],
                category_title=row[2],
                score=row[3],
                total=row[4],
                percentage=float(row[5]),
                time_spent_seconds=row[6],
                created_at=row[7]
            )
            for row in recent_rows
        ]
        
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [QUIZ] Retrieved stats for user {user_id}: {total_quizzes} quizzes, {average_score:.1f}% avg")
        
        return QuizStatsResponse(
            total_quizzes=total_quizzes,
            average_score=average_score,
            best_category=best_category,
            best_category_title=best_category_title,
            best_category_percentage=best_category_percentage,
            recent_results=recent_results
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [QUIZ] Failed to get stats: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get quiz stats: {str(e)}"
        )


@app.get("/api/quiz/history", tags=["Quiz"])
async def get_quiz_history(
    page: int = 1,
    per_page: int = 20,
    authorization: Optional[str] = Header(None)
):
    """
    Get paginated quiz history for the authenticated user.
    
    Returns all quiz results with pagination metadata.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Validate pagination params
        if page < 1:
            page = 1
        if per_page < 1 or per_page > 100:
            per_page = 20
        
        offset = (page - 1) * per_page
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get total count
        cursor.execute("""
            SELECT COUNT(*) FROM quiz_results WHERE user_id = %s
        """, (user_id,))
        total_count = cursor.fetchone()[0]
        
        # Get paginated results
        cursor.execute("""
            SELECT id, category_id, category_title, score, total, percentage, time_spent_seconds, created_at
            FROM quiz_results
            WHERE user_id = %s
            ORDER BY created_at DESC
            LIMIT %s OFFSET %s
        """, (user_id, per_page, offset))
        
        rows = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        results = [
            QuizResultResponse(
                id=str(row[0]),
                category_id=row[1],
                category_title=row[2],
                score=row[3],
                total=row[4],
                percentage=float(row[5]),
                time_spent_seconds=row[6],
                created_at=row[7]
            )
            for row in rows
        ]
        
        total_pages = (total_count + per_page - 1) // per_page  # Ceiling division
        
        return {
            "results": results,
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total_count": total_count,
                "total_pages": total_pages,
                "has_next": page < total_pages,
                "has_prev": page > 1
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [QUIZ] Failed to get history: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get quiz history: {str(e)}"
        )


@app.get("/api/quiz/weak-areas", tags=["Quiz"])
async def get_weak_areas(
    authorization: Optional[str] = Header(None)
):
    """
    Analyze the user's quiz history to identify weak areas (categories with low scores).
    
    Returns categories where the user is struggling, sorted by priority (lowest scores first).
    Only considers categories with at least 1 quiz attempt in the last 30 days.
    """
    try:
        user_id = await verify_user(authorization)
        
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get average scores per category from recent quizzes (last 30 days)
        cursor.execute("""
            SELECT 
                category_id,
                category_title,
                AVG(percentage) as avg_score,
                COUNT(*) as attempt_count,
                MAX(created_at) as last_attempt
            FROM quiz_results
            WHERE user_id = %s 
            AND created_at >= NOW() - INTERVAL '30 days'
            GROUP BY category_id, category_title
            ORDER BY avg_score ASC
        """, (user_id,))
        
        rows = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        weak_areas = []
        for row in rows:
            category_id, category_title, avg_score, attempt_count, last_attempt = row
            
            # Consider it a "weak area" if average score is below 70%
            if avg_score < 70:
                weak_areas.append({
                    "category_id": category_id,
                    "category_title": category_title,
                    "avg_score": round(float(avg_score), 1),
                    "attempt_count": attempt_count,
                    "last_attempt": last_attempt.isoformat() if last_attempt else None,
                    "priority": "high" if avg_score < 50 else "medium"
                })
        
        return {
            "weak_areas": weak_areas,
            "has_weak_areas": len(weak_areas) > 0,
            "recommendation": weak_areas[0] if weak_areas else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [QUIZ] Failed to get weak areas: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get weak areas: {str(e)}"
        )


# ==========================================
# Game Results Endpoints
# ==========================================

@app.post("/api/games/results", response_model=GameResultResponse, tags=["Games"])
async def save_game_result(
    request: GameResultCreate,
    authorization: Optional[str] = Header(None)
):
    """
    Save a game result for the authenticated user.
    
    Requires authentication.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Insert game result
        cursor.execute("""
            INSERT INTO game_results (
                user_id, game_type, mode, category_id, category_title,
                difficulty, score, moves, pairs, time_seconds, stars
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id, created_at
        """, (
            user_id,
            request.game_type,
            request.mode,
            request.category_id,
            request.category_title,
            request.difficulty,
            request.score,
            request.moves,
            request.pairs,
            request.time_seconds,
            request.stars
        ))
        
        result = cursor.fetchone()
        result_id = result[0]
        created_at = result[1]
        
        conn.commit()
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [GAME] Saved {request.game_type} result for user {user_id}: {request.score} pts, {request.stars} stars")
        
        return GameResultResponse(
            id=str(result_id),
            game_type=request.game_type,
            mode=request.mode,
            category_id=request.category_id,
            category_title=request.category_title,
            difficulty=request.difficulty,
            score=request.score,
            moves=request.moves,
            pairs=request.pairs,
            time_seconds=request.time_seconds,
            stars=request.stars,
            created_at=created_at
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [GAME] Failed to save result: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save game result: {str(e)}"
        )


@app.get("/api/games/stats", response_model=GameStatsResponse, tags=["Games"])
async def get_game_stats(
    authorization: Optional[str] = Header(None)
):
    """
    Get game statistics for the authenticated user.
    
    Returns total games, average score, best category, and recent results.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        # Get total games and averages
        cursor.execute("""
            SELECT 
                COUNT(*) as total_games,
                COALESCE(AVG(score), 0) as average_score,
                COALESCE(AVG(stars), 0) as average_stars
            FROM game_results
            WHERE user_id = %s
        """, (user_id,))
        
        stats = cursor.fetchone()
        total_games = stats[0]
        average_score = float(stats[1])
        average_stars = float(stats[2])
        
        # Get best category (highest average score, min 1 attempt)
        cursor.execute("""
            SELECT 
                category_id,
                category_title,
                AVG(score) as avg_score
            FROM game_results
            WHERE user_id = %s
            GROUP BY category_id, category_title
            HAVING COUNT(*) >= 1
            ORDER BY avg_score DESC
            LIMIT 1
        """, (user_id,))
        
        best_category_row = cursor.fetchone()
        best_category = None
        best_category_title = None
        best_category_score = None
        
        if best_category_row:
            best_category = best_category_row[0]
            best_category_title = best_category_row[1]
            best_category_score = float(best_category_row[2])
        
        # Get recent results (last 10)
        cursor.execute("""
            SELECT id, game_type, mode, category_id, category_title, difficulty,
                   score, moves, pairs, time_seconds, stars, created_at
            FROM game_results
            WHERE user_id = %s
            ORDER BY created_at DESC
            LIMIT 10
        """, (user_id,))
        
        recent_rows = cursor.fetchall()
        recent_results = [
            GameResultResponse(
                id=str(row[0]),
                game_type=row[1],
                mode=row[2],
                category_id=row[3],
                category_title=row[4],
                difficulty=row[5],
                score=row[6],
                moves=row[7],
                pairs=row[8],
                time_seconds=row[9],
                stars=row[10],
                created_at=row[11]
            )
            for row in recent_rows
        ]
        
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [GAME] Retrieved stats for user {user_id}: {total_games} games, {average_score:.0f} avg score")
        
        return GameStatsResponse(
            total_games=total_games,
            average_score=average_score,
            average_stars=average_stars,
            best_category=best_category,
            best_category_title=best_category_title,
            best_category_score=best_category_score,
            recent_results=recent_results
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [GAME] Failed to get stats: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get game stats: {str(e)}"
        )


@app.get("/api/games/history", tags=["Games"])
async def get_game_history(
    authorization: Optional[str] = Header(None),
    page: int = Query(1, ge=1, description="Page number"),
    per_page: int = Query(10, ge=1, le=50, description="Results per page"),
    game_type: Optional[str] = Query(None, description="Filter by game type")
):
    """
    Get paginated game history for the authenticated user.
    
    Optionally filter by game type.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        offset = (page - 1) * per_page
        
        # Build query with optional game_type filter
        if game_type:
            cursor.execute("""
                SELECT COUNT(*) FROM game_results WHERE user_id = %s AND game_type = %s
            """, (user_id, game_type))
        else:
            cursor.execute("""
                SELECT COUNT(*) FROM game_results WHERE user_id = %s
            """, (user_id,))
        
        total_count = cursor.fetchone()[0]
        
        # Get paginated results
        if game_type:
            cursor.execute("""
                SELECT id, game_type, mode, category_id, category_title, difficulty,
                       score, moves, pairs, time_seconds, stars, created_at
                FROM game_results
                WHERE user_id = %s AND game_type = %s
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            """, (user_id, game_type, per_page, offset))
        else:
            cursor.execute("""
                SELECT id, game_type, mode, category_id, category_title, difficulty,
                       score, moves, pairs, time_seconds, stars, created_at
                FROM game_results
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            """, (user_id, per_page, offset))
        
        rows = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        results = [
            GameResultResponse(
                id=str(row[0]),
                game_type=row[1],
                mode=row[2],
                category_id=row[3],
                category_title=row[4],
                difficulty=row[5],
                score=row[6],
                moves=row[7],
                pairs=row[8],
                time_seconds=row[9],
                stars=row[10],
                created_at=row[11]
            )
            for row in rows
        ]
        
        total_pages = (total_count + per_page - 1) // per_page
        
        return {
            "results": results,
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total_count": total_count,
                "total_pages": total_pages,
                "has_next": page < total_pages,
                "has_prev": page > 1
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [GAME] Failed to get history: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get game history: {str(e)}"
        )


# ==========================================
# Usage Tracking Endpoints (Freemium Limits)
# ==========================================

from .models import UsageResponse, UsageIncrementResponse, SubscriptionStatusResponse
from datetime import date, timezone, timedelta

# Guam timezone (ChST = UTC+10)
GUAM_TIMEZONE = timezone(timedelta(hours=10))

def get_guam_date() -> date:
    """Get the current date in Guam timezone (ChST, UTC+10)."""
    return datetime.now(GUAM_TIMEZONE).date()

# Default daily limits for free users
FREE_TIER_LIMITS = {
    "chat": 8,
    "game": 10,
    "quiz": 5
}

# Holiday Promo Period Configuration
# Set FREE_PROMO_PERIOD=true in environment to give everyone unlimited access
FREE_PROMO_ACTIVE = os.getenv("FREE_PROMO_PERIOD", "false").lower() == "true"
FREE_PROMO_END_DATE = os.getenv("FREE_PROMO_END_DATE", "2026-01-06")  # Three Kings Day

def is_promo_active() -> bool:
    """Check if the free promo period is currently active."""
    if not FREE_PROMO_ACTIVE:
        return False
    try:
        end_date = datetime.strptime(FREE_PROMO_END_DATE, "%Y-%m-%d").date()
        today = get_guam_date()
        return today <= end_date
    except ValueError:
        return FREE_PROMO_ACTIVE  # Fall back to env var if date parsing fails


@app.get("/api/promo/status", tags=["Promo"])
async def get_promo_status():
    """
    Check if a promotional period is currently active and get current theme.
    
    Returns promo status, end date, and theme for frontend display.
    No authentication required - public endpoint.
    
    Reads from database settings (admin-configurable) with fallback to env vars.
    """
    # Try database first, fall back to env var for backwards compatibility
    promo_active = is_promo_active_from_db() if get_site_setting("promo_enabled") else is_promo_active()
    promo_end_date = get_site_setting("promo_end_date", FREE_PROMO_END_DATE)
    promo_title = get_site_setting("promo_title", "🎄 Holiday Special: Unlimited access!")
    current_theme = get_site_setting("theme", "default")
    
    return {
        "active": promo_active,
        "end_date": promo_end_date if promo_active else None,
        "message": promo_title if promo_active else None,
        "theme": current_theme
    }


@app.get("/api/usage/today", response_model=UsageResponse, tags=["Usage"])
async def get_today_usage(
    authorization: Optional[str] = Header(None)
):
    """
    Get the current user's usage for today.
    
    Returns current counts and limits for chat, games, and quizzes.
    Premium users have unlimited access (limits set to -1).
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Check if user is premium OR if promo period is active (from DB or env)
        is_premium = is_promo_active_from_db() if get_site_setting("promo_enabled") else is_promo_active()
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        today = get_guam_date()  # Use Guam timezone for daily reset
        
        # Get or create usage record for today
        cursor.execute("""
            SELECT chat_count, game_count, quiz_count
            FROM user_daily_usage
            WHERE user_id = %s AND usage_date = %s
        """, (user_id, today))
        
        row = cursor.fetchone()
        
        if row:
            chat_count, game_count, quiz_count = row
        else:
            chat_count, game_count, quiz_count = 0, 0, 0
        
        cursor.close()
        conn.close()
        
        # Premium users have unlimited (-1 means no limit)
        if is_premium:
            return UsageResponse(
                chat_count=chat_count,
                game_count=game_count,
                quiz_count=quiz_count,
                chat_limit=-1,
                game_limit=-1,
                quiz_limit=-1,
                is_premium=True
            )
        
        return UsageResponse(
            chat_count=chat_count,
            game_count=game_count,
            quiz_count=quiz_count,
            chat_limit=FREE_TIER_LIMITS["chat"],
            game_limit=FREE_TIER_LIMITS["game"],
            quiz_limit=FREE_TIER_LIMITS["quiz"],
            is_premium=False
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [USAGE] Failed to get usage: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get usage: {str(e)}"
        )


@app.post("/api/usage/increment/{usage_type}", response_model=UsageIncrementResponse, tags=["Usage"])
async def increment_usage(
    usage_type: str,
    authorization: Optional[str] = Header(None)
):
    """
    Increment usage counter for a specific feature type.
    
    Args:
        usage_type: One of 'chat', 'game', or 'quiz'
    
    Returns:
        - success: True if the action was allowed (under limit or premium)
        - new_count: The new usage count after increment
        - remaining: How many uses left today (-1 for unlimited)
    
    For free users, this will fail if they've hit their daily limit.
    """
    if usage_type not in ["chat", "game", "quiz"]:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid usage type: {usage_type}. Must be 'chat', 'game', or 'quiz'"
        )
    
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        # Check if user is premium OR if promo period is active (from DB or env)
        is_premium = is_promo_active_from_db() if get_site_setting("promo_enabled") else is_promo_active()
        
        # Get database connection
        conn = conversations.get_db_connection()
        cursor = conn.cursor()
        
        today = get_guam_date()  # Use Guam timezone for daily reset
        column_name = f"{usage_type}_count"
        limit = FREE_TIER_LIMITS[usage_type]
        
        # Get current usage
        cursor.execute("""
            SELECT chat_count, game_count, quiz_count
            FROM user_daily_usage
            WHERE user_id = %s AND usage_date = %s
        """, (user_id, today))
        
        row = cursor.fetchone()
        
        if row:
            current_counts = {"chat": row[0], "game": row[1], "quiz": row[2]}
            current_count = current_counts[usage_type]
        else:
            current_count = 0
        
        # Check if under limit (premium users always pass)
        if not is_premium and current_count >= limit:
            cursor.close()
            conn.close()
            return UsageIncrementResponse(
                success=False,
                new_count=current_count,
                limit=limit,
                remaining=0,
                is_premium=False
            )
        
        # Increment the counter using UPSERT
        cursor.execute(f"""
            INSERT INTO user_daily_usage (user_id, usage_date, {column_name})
            VALUES (%s, %s, 1)
            ON CONFLICT (user_id, usage_date)
            DO UPDATE SET 
                {column_name} = user_daily_usage.{column_name} + 1,
                updated_at = now()
            RETURNING {column_name}
        """, (user_id, today))
        
        new_count = cursor.fetchone()[0]
        
        conn.commit()
        cursor.close()
        conn.close()
        
        # Calculate remaining (-1 for unlimited)
        if is_premium:
            remaining = -1
        else:
            remaining = max(0, limit - new_count)
        
        return UsageIncrementResponse(
            success=True,
            new_count=new_count,
            limit=limit if not is_premium else -1,
            remaining=remaining,
            is_premium=is_premium
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [USAGE] Failed to increment usage: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to increment usage: {str(e)}"
        )


# ==========================================
# Streak Tracking Endpoint
# ==========================================

@app.get("/api/streaks/me", tags=["Streaks"])
async def get_user_streak(
    authorization: Optional[str] = Header(None)
):
    """
    Get the current user's learning streak.
    
    A streak is the number of consecutive days with at least one activity
    (chat message, game, or quiz). Uses Guam timezone for day boundaries.
    
    Returns:
    - current_streak: Number of consecutive days (including today if active)
    - longest_streak: User's all-time best streak
    - is_today_active: Whether the user has done any activity today
    - today_activities: Breakdown of today's activities
    - last_activity_date: The most recent day with activity
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        today = get_guam_date()
        
        # Get all activity days for this user (ordered newest first)
        # Includes: chat, games, quizzes, AND learning activities (lessons/flashcards)
        cursor.execute("""
            SELECT DISTINCT activity_date FROM (
                -- Chat, games, quizzes from daily usage
                SELECT usage_date as activity_date
            FROM user_daily_usage
            WHERE user_id = %s
            AND (chat_count > 0 OR game_count > 0 OR quiz_count > 0)
                
                UNION
                
                -- Learning activities (lesson started or completed)
                SELECT DATE(started_at AT TIME ZONE 'Pacific/Guam') as activity_date
                FROM user_topic_progress
                WHERE user_id = %s AND started_at IS NOT NULL
                
                UNION
                
                -- Learning activities (lesson completed)
                SELECT DATE(completed_at AT TIME ZONE 'Pacific/Guam') as activity_date
                FROM user_topic_progress
                WHERE user_id = %s AND completed_at IS NOT NULL
            ) combined_activities
            ORDER BY activity_date DESC
        """, (user_id, user_id, user_id))
        
        activity_days = [(row[0],) for row in cursor.fetchall()]  # Convert to same format
        
        # Calculate current streak
        current_streak = 0
        last_activity_date = None
        
        if activity_days:
            # Check if there's activity today or yesterday to have an active streak
            check_date = today
            
            for row in activity_days:
                activity_date = row[0]
                
                if activity_date == check_date:
                    current_streak += 1
                    check_date = check_date - timedelta(days=1)
                    if last_activity_date is None:
                        last_activity_date = activity_date
                elif activity_date == check_date - timedelta(days=1):
                    # Allow for yesterday to start a streak (if today not yet active)
                    if current_streak == 0:
                        check_date = activity_date
                        current_streak = 1
                        last_activity_date = activity_date
                        check_date = check_date - timedelta(days=1)
                    else:
                        break
                else:
                    break
        
        # Calculate longest streak (scan through all activity days)
        longest_streak = current_streak
        if activity_days:
            temp_streak = 1
            prev_date = activity_days[0][0]
            
            for row in activity_days[1:]:
                activity_date = row[0]
                if prev_date - activity_date == timedelta(days=1):
                    temp_streak += 1
                    longest_streak = max(longest_streak, temp_streak)
                else:
                    temp_streak = 1
                prev_date = activity_date
        
        # Get today's activity breakdown
        today_chat = 0
        today_games = 0
        today_quizzes = 0
        today_learning = 0
        is_today_active = False
        
        cursor.execute("""
            SELECT chat_count, game_count, quiz_count
            FROM user_daily_usage
            WHERE user_id = %s AND usage_date = %s
        """, (user_id, today))
        
        today_row = cursor.fetchone()
        if today_row:
            today_chat = today_row[0] or 0
            today_games = today_row[1] or 0
            today_quizzes = today_row[2] or 0
        
        # Check for learning activities today
        cursor.execute("""
            SELECT COUNT(*) FROM user_topic_progress
            WHERE user_id = %s 
            AND (
                DATE(started_at AT TIME ZONE 'Pacific/Guam') = %s
                OR DATE(completed_at AT TIME ZONE 'Pacific/Guam') = %s
            )
        """, (user_id, today, today))
        
        learning_row = cursor.fetchone()
        if learning_row:
            today_learning = learning_row[0] or 0
        
        is_today_active = (today_chat + today_games + today_quizzes + today_learning) > 0
        
        conn.close()
        
        return {
            "current_streak": current_streak,
            "longest_streak": longest_streak,
            "is_today_active": is_today_active,
            "today_activities": {
                "chat": today_chat,
                "games": today_games,
                "quizzes": today_quizzes,
                "learning": today_learning
            },
            "last_activity_date": str(last_activity_date) if last_activity_date else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [STREAKS] Failed to get streak: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get streak: {str(e)}"
        )


# ==========================================
# Unified Homepage Data Endpoint
# ==========================================

def _fetch_streak_data(user_id: str, db_url: str):
    """Fetch streak data for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    today = get_guam_date()
    
    # Get all activity days for this user
    cursor.execute("""
        SELECT DISTINCT activity_date FROM (
            SELECT usage_date as activity_date
            FROM user_daily_usage
            WHERE user_id = %s
            AND (chat_count > 0 OR game_count > 0 OR quiz_count > 0)
            UNION
            SELECT DATE(started_at AT TIME ZONE 'Pacific/Guam') as activity_date
            FROM user_topic_progress
            WHERE user_id = %s AND started_at IS NOT NULL
            UNION
            SELECT DATE(completed_at AT TIME ZONE 'Pacific/Guam') as activity_date
            FROM user_topic_progress
            WHERE user_id = %s AND completed_at IS NOT NULL
        ) combined_activities
        ORDER BY activity_date DESC
    """, (user_id, user_id, user_id))
    
    activity_days = [(row[0],) for row in cursor.fetchall()]
    
    current_streak = 0
    last_activity_date = None
    
    if activity_days:
        check_date = today
        for row in activity_days:
            activity_date = row[0]
            if activity_date == check_date:
                current_streak += 1
                check_date = check_date - timedelta(days=1)
            elif activity_date == check_date - timedelta(days=1):
                current_streak += 1
                check_date = activity_date - timedelta(days=1)
            else:
                break
            last_activity_date = activity_date
    
    # Get today's activity breakdown
    cursor.execute("""
        SELECT chat_count, game_count, quiz_count
        FROM user_daily_usage
        WHERE user_id = %s AND usage_date = %s
    """, (user_id, today))
    
    today_row = cursor.fetchone()
    today_chat = today_row[0] or 0 if today_row else 0
    today_games = today_row[1] or 0 if today_row else 0
    today_quizzes = today_row[2] or 0 if today_row else 0
    
    cursor.execute("""
        SELECT COUNT(*) FROM user_topic_progress
        WHERE user_id = %s 
        AND (
            DATE(started_at AT TIME ZONE 'Pacific/Guam') = %s
            OR DATE(completed_at AT TIME ZONE 'Pacific/Guam') = %s
        )
    """, (user_id, today, today))
    
    learning_row = cursor.fetchone()
    today_learning = learning_row[0] or 0 if learning_row else 0
    is_today_active = (today_chat + today_games + today_quizzes + today_learning) > 0
    
    cursor.close()
    conn.close()
    
    return {
        "current_streak": current_streak,
        "longest_streak": current_streak,  # Simplified for homepage
        "is_today_active": is_today_active,
        "today_activities": {
            "chat": today_chat,
            "games": today_games,
            "quizzes": today_quizzes,
            "learning": today_learning
        },
        "last_activity_date": str(last_activity_date) if last_activity_date else None
    }


def _fetch_xp_data(user_id: str, db_url: str):
    """Fetch XP data for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    today = get_guam_date()
    
    cursor.execute("""
        SELECT total_xp, level, daily_goal_minutes, today_minutes, goal_date
        FROM user_xp WHERE user_id = %s
    """, (user_id,))
    
    row = cursor.fetchone()
    
    if row:
        total_xp, level, daily_goal_minutes, today_minutes, goal_date = row
        if goal_date != today:
            today_minutes = 0
    else:
        total_xp, level, daily_goal_minutes, today_minutes = 0, 1, 10, 0
    
    cursor.close()
    conn.close()
    
    level = calculate_level(total_xp)
    xp_for_current = LEVEL_THRESHOLDS[level - 1] if level > 1 else 0
    xp_for_next = get_xp_for_next_level(level)
    
    xp_in_level = total_xp - xp_for_current
    xp_needed = xp_for_next - xp_for_current
    xp_progress = min(100, int((xp_in_level / xp_needed) * 100)) if xp_needed > 0 else 100
    
    return {
        "total_xp": total_xp,
        "level": level,
        "xp_for_current_level": xp_for_current,
        "xp_for_next_level": xp_for_next,
        "xp_progress": xp_progress,
        "daily_goal_minutes": daily_goal_minutes,
        "today_minutes": today_minutes,
        "daily_goal_complete": today_minutes >= daily_goal_minutes if daily_goal_minutes > 0 else True
    }


def _fetch_quiz_stats(user_id: str, db_url: str):
    """Fetch quiz stats for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT COUNT(*) as total_quizzes, COALESCE(AVG(percentage), 0) as average_score
        FROM quiz_results WHERE user_id = %s
    """, (user_id,))
    
    stats = cursor.fetchone()
    total_quizzes = stats[0]
    average_score = float(stats[1])
    
    cursor.close()
    conn.close()
    
    return {
        "total_quizzes": total_quizzes,
        "average_score": round(average_score, 1)
    }


def _fetch_game_stats(user_id: str, db_url: str):
    """Fetch game stats for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT COUNT(*) as total_games, COALESCE(AVG(score), 0) as average_score
        FROM game_results WHERE user_id = %s
    """, (user_id,))
    
    stats = cursor.fetchone()
    
    cursor.close()
    conn.close()
    
    return {
        "total_games": stats[0],
        "average_score": round(float(stats[1]), 1)
    }


def _fetch_weak_areas(user_id: str, db_url: str):
    """Fetch weak areas for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT category_id, category_title, AVG(percentage) as avg_score, COUNT(*) as attempt_count
        FROM quiz_results
        WHERE user_id = %s AND created_at >= NOW() - INTERVAL '30 days'
        GROUP BY category_id, category_title
        ORDER BY avg_score ASC
    """, (user_id,))
    
    rows = cursor.fetchall()
    cursor.close()
    conn.close()
    
    weak_areas = []
    for row in rows:
        category_id, category_title, avg_score, attempt_count = row
        if avg_score < 70:
            weak_areas.append({
                "category_id": category_id,
                "category_title": category_title,
                "avg_score": round(float(avg_score), 1),
                "attempt_count": attempt_count,
                "priority": "high" if avg_score < 50 else "medium"
            })
    
    return {
        "weak_areas": weak_areas,
        "has_weak_areas": len(weak_areas) > 0,
        "recommendation": weak_areas[0] if weak_areas else None
    }


def _fetch_sr_summary(user_id: str, db_url: str):
    """Fetch spaced repetition summary for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            COUNT(*) as total_cards,
            COUNT(*) FILTER (WHERE next_review IS NULL OR next_review <= NOW()) as due_today,
            COUNT(*) FILTER (WHERE easiness_factor > 2.3 AND interval > 30) as mastered,
            COUNT(*) FILTER (WHERE interval < 7) as learning
        FROM spaced_repetition
        WHERE user_id = %s
    """, (user_id,))
    
    row = cursor.fetchone()
    cursor.close()
    conn.close()
    
    if row:
        return {
            "total_cards": row[0],
            "due_today": row[1],
            "mastered": row[2],
            "learning": row[3],
            "has_cards": row[0] > 0
        }
    
    return {"total_cards": 0, "due_today": 0, "mastered": 0, "learning": 0, "has_cards": False}


def _fetch_learning_recommended(user_id: str, db_url: str):
    """Fetch recommended topic for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT topic_id, started_at, completed_at, best_quiz_score, flashcards_viewed, last_activity_at
        FROM user_topic_progress WHERE user_id = %s
    """, (user_id,))
    
    progress_rows = cursor.fetchall()
    cursor.close()
    conn.close()
    
    # Build progress map
    progress_map = {}
    for row in progress_rows:
        progress_map[row[0]] = {
            "topic_id": row[0],
            "started_at": row[1].isoformat() if row[1] else None,
            "completed_at": row[2].isoformat() if row[2] else None,
            "best_quiz_score": row[3],
            "flashcards_viewed": row[4] or 0,
            "last_activity_at": row[5].isoformat() if row[5] else None,
        }
    
    # Count completions per level
    beginner_completed = sum(1 for t in BEGINNER_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    intermediate_completed = sum(1 for t in INTERMEDIATE_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    advanced_completed = sum(1 for t in ADVANCED_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    total_completed = beginner_completed + intermediate_completed + advanced_completed
    
    beginner_complete = beginner_completed == len(BEGINNER_PATH)
    intermediate_complete = intermediate_completed == len(INTERMEDIATE_PATH)
    
    recommendation_type = "start"
    recommended_topic = None
    topic_progress = None
    message = ""
    
    # Check each path in order
    paths = [
        (BEGINNER_PATH, beginner_complete, "intermediate"),
        (INTERMEDIATE_PATH, intermediate_complete, "advanced"),
        (ADVANCED_PATH, True, None)
    ]
    
    for path, path_complete, next_level in paths:
        if path_complete:
            continue
        for topic in path:
            topic_id = topic["id"]
            progress = progress_map.get(topic_id)
            
            if not progress:
                recommendation_type = "start" if total_completed == 0 else "next"
                recommended_topic = topic
                topic_progress = {"topic_id": topic_id, "started_at": None, "completed_at": None, "best_quiz_score": None, "flashcards_viewed": 0}
                message = f"Start your Chamorro journey with {topic['title']}!" if total_completed == 0 else f"Great progress! Ready for {topic['title']}?"
                break
            elif not progress.get("completed_at"):
                recommendation_type = "continue"
                recommended_topic = topic
                topic_progress = progress
                message = f"Continue where you left off in {topic['title']}"
                break
        if recommended_topic:
            break
    
    if not recommended_topic:
        recommendation_type = "complete"
        message = "Congratulations! You've completed all topics! Review any topic to keep your skills sharp."
    
    return {
        "recommendation_type": recommendation_type,
        "topic": recommended_topic,
        "progress": topic_progress,
        "completed_topics": total_completed,
        "total_topics": len(ALL_TOPICS),
        "message": message
    }


def _fetch_all_progress(user_id: str, db_url: str):
    """Fetch all learning progress for user (runs in thread pool)."""
    import psycopg2
    conn = psycopg2.connect(db_url)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT topic_id, started_at, completed_at, best_quiz_score, flashcards_viewed, last_activity_at
        FROM user_topic_progress WHERE user_id = %s
    """, (user_id,))
    
    progress_rows = cursor.fetchall()
    cursor.close()
    conn.close()
    
    progress_map = {}
    for row in progress_rows:
        progress_map[row[0]] = {
            "topic_id": row[0],
            "started_at": row[1].isoformat() if row[1] else None,
            "completed_at": row[2].isoformat() if row[2] else None,
            "best_quiz_score": row[3],
            "flashcards_viewed": row[4] or 0,
            "last_activity_at": row[5].isoformat() if row[5] else None,
        }
    
    topics_with_progress = []
    for topic in ALL_TOPICS:
        topic_id = topic["id"]
        progress = progress_map.get(topic_id, {
            "topic_id": topic_id, "started_at": None, "completed_at": None,
            "best_quiz_score": None, "flashcards_viewed": 0, "last_activity_at": None,
        })
        topics_with_progress.append({"topic": topic, "progress": progress})
    
    beginner_completed = sum(1 for t in BEGINNER_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    intermediate_completed = sum(1 for t in INTERMEDIATE_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    advanced_completed = sum(1 for t in ADVANCED_PATH if progress_map.get(t["id"], {}).get("completed_at"))
    
    return {
        "topics": topics_with_progress,
        "summary": {
            "beginner": {"completed": beginner_completed, "total": len(BEGINNER_PATH)},
            "intermediate": {"completed": intermediate_completed, "total": len(INTERMEDIATE_PATH)},
            "advanced": {"completed": advanced_completed, "total": len(ADVANCED_PATH)},
            "total_completed": beginner_completed + intermediate_completed + advanced_completed,
            "total_topics": len(ALL_TOPICS)
        }
    }


@app.get("/api/homepage/data", tags=["Homepage"])
async def get_homepage_data(
    authorization: Optional[str] = Header(None)
):
    """
    Unified endpoint that fetches all data needed for the homepage in one call.
    
    This runs all queries in parallel for maximum performance, reducing
    homepage load from 8+ API calls down to 1.
    
    Returns: streak, xp, quiz_stats, game_stats, weak_areas, sr_summary, 
             recommended_topic, all_progress
    """
    try:
        user_id = await verify_user(authorization)
        db_url = os.getenv("DATABASE_URL")
        
        # Run all data fetching in parallel using asyncio.gather
        # Each function runs its DB queries synchronously but we run them concurrently in thread pool
        loop = asyncio.get_event_loop()
        
        results = await asyncio.gather(
            loop.run_in_executor(None, _fetch_streak_data, user_id, db_url),
            loop.run_in_executor(None, _fetch_xp_data, user_id, db_url),
            loop.run_in_executor(None, _fetch_quiz_stats, user_id, db_url),
            loop.run_in_executor(None, _fetch_game_stats, user_id, db_url),
            loop.run_in_executor(None, _fetch_weak_areas, user_id, db_url),
            loop.run_in_executor(None, _fetch_sr_summary, user_id, db_url),
            loop.run_in_executor(None, _fetch_learning_recommended, user_id, db_url),
            loop.run_in_executor(None, _fetch_all_progress, user_id, db_url),
            return_exceptions=True
        )
        
        # Handle any exceptions
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"❌ [HOMEPAGE] Task {i} failed: {result}")
                results[i] = None
        
        streak, xp, quiz_stats, game_stats, weak_areas, sr_summary, recommended, all_progress = results
        
        return {
            "streak": streak,
            "xp": xp,
            "quiz_stats": quiz_stats,
            "game_stats": game_stats,
            "weak_areas": weak_areas,
            "sr_summary": sr_summary,
            "recommended": recommended,
            "all_progress": all_progress
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [HOMEPAGE] Failed to get homepage data: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get homepage data: {str(e)}"
        )


# ==========================================
# XP & Leveling System
# ==========================================

# XP values for different activities
XP_VALUES = {
    "flashcard_complete": 10,      # Completing a flashcard deck
    "quiz_complete": 25,           # Completing a quiz
    "quiz_bonus_90": 10,           # Bonus for 90%+ on quiz
    "game_complete": 5,            # Completing a game
    "chat_message": 2,             # Sending a chat message
    "topic_complete": 50,          # Completing a learning topic
    "streak_bonus": 15,            # Daily streak maintained
    "daily_goal_complete": 20,     # Completing daily goal
}

# Level thresholds (cumulative XP needed for each level)
LEVEL_THRESHOLDS = [
    0,      # Level 1: 0 XP
    100,    # Level 2: 100 XP
    250,    # Level 3: 250 XP
    500,    # Level 4: 500 XP
    850,    # Level 5: 850 XP
    1300,   # Level 6: 1300 XP
    1900,   # Level 7: 1900 XP
    2600,   # Level 8: 2600 XP
    3500,   # Level 9: 3500 XP
    4600,   # Level 10: 4600 XP
    6000,   # Level 11: 6000 XP
    7700,   # Level 12: 7700 XP
    9700,   # Level 13: 9700 XP
    12000,  # Level 14: 12000 XP
    15000,  # Level 15: 15000 XP (max for now)
]

def calculate_level(total_xp: int) -> int:
    """Calculate the user's level based on total XP."""
    for level, threshold in enumerate(LEVEL_THRESHOLDS):
        if total_xp < threshold:
            return level  # Return the previous level
    return len(LEVEL_THRESHOLDS)  # Max level


def get_xp_for_next_level(level: int) -> int:
    """Get the XP threshold for the next level."""
    if level >= len(LEVEL_THRESHOLDS):
        return LEVEL_THRESHOLDS[-1]  # Already at max
    return LEVEL_THRESHOLDS[level]


@app.get("/api/xp/me", tags=["XP"])
async def get_user_xp(
    authorization: Optional[str] = Header(None)
):
    """
    Get the current user's XP, level, and progress.
    
    Returns:
    - total_xp: Total XP earned
    - level: Current level (1-15)
    - xp_for_current_level: XP threshold for current level
    - xp_for_next_level: XP threshold for next level
    - xp_progress: Progress towards next level (0-100%)
    - daily_goal_minutes: User's daily goal setting
    - today_minutes: Minutes spent today
    - daily_goal_complete: Whether today's goal is met
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        today = get_guam_date()
        
        # Get or create user XP record
        cursor.execute("""
            SELECT total_xp, level, daily_goal_minutes, today_minutes, goal_date
            FROM user_xp WHERE user_id = %s
        """, (user_id,))
        
        row = cursor.fetchone()
        
        if row:
            total_xp, level, daily_goal_minutes, today_minutes, goal_date = row
            # Reset today_minutes if it's a new day
            if goal_date != today:
                today_minutes = 0
        else:
            # Create new XP record for user
            cursor.execute("""
                INSERT INTO user_xp (user_id, total_xp, level, daily_goal_minutes, today_minutes, goal_date)
                VALUES (%s, 0, 1, 10, 0, %s)
                RETURNING total_xp, level, daily_goal_minutes, today_minutes
            """, (user_id, today))
            total_xp, level, daily_goal_minutes, today_minutes = cursor.fetchone()
            conn.commit()
        
        cursor.close()
        conn.close()
        
        # Calculate level and progress
        level = calculate_level(total_xp)
        xp_for_current = LEVEL_THRESHOLDS[level - 1] if level > 1 else 0
        xp_for_next = get_xp_for_next_level(level)
        
        xp_in_level = total_xp - xp_for_current
        xp_needed = xp_for_next - xp_for_current
        xp_progress = min(100, int((xp_in_level / xp_needed) * 100)) if xp_needed > 0 else 100
        
        return {
            "total_xp": total_xp,
            "level": level,
            "xp_for_current_level": xp_for_current,
            "xp_for_next_level": xp_for_next,
            "xp_progress": xp_progress,
            "daily_goal_minutes": daily_goal_minutes,
            "today_minutes": today_minutes,
            "daily_goal_complete": today_minutes >= daily_goal_minutes if daily_goal_minutes > 0 else True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [XP] Failed to get XP: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get XP: {str(e)}"
        )


@app.post("/api/xp/award", tags=["XP"])
async def award_xp(
    request: dict,
    authorization: Optional[str] = Header(None)
):
    """
    Award XP to a user for an activity.
    
    Request body:
    - activity_type: One of 'flashcard_complete', 'quiz_complete', 'game_complete', 
                     'chat_message', 'topic_complete', 'streak_bonus', 'daily_goal_complete'
    - activity_id: Optional reference (topic_id, quiz_id, etc.)
    - quiz_score: Optional score (0-100) for quiz bonus calculation
    - minutes_spent: Optional time spent for daily goal tracking
    
    Returns:
    - xp_earned: Amount of XP earned
    - total_xp: New total XP
    - level: Current level
    - level_up: Whether user leveled up from this award
    - new_level: New level if leveled up
    """
    try:
        user_id = await verify_user(authorization)
        
        activity_type = request.get("activity_type")
        activity_id = request.get("activity_id")
        quiz_score = request.get("quiz_score", 0)
        minutes_spent = request.get("minutes_spent", 0)
        
        if not activity_type or activity_type not in XP_VALUES:
            raise HTTPException(status_code=400, detail=f"Invalid activity_type: {activity_type}")
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        today = get_guam_date()
        
        # Calculate XP to award
        xp_earned = XP_VALUES[activity_type]
        description = activity_type.replace("_", " ").title()
        
        # Quiz bonus for 90%+ score
        if activity_type == "quiz_complete" and quiz_score >= 90:
            xp_earned += XP_VALUES["quiz_bonus_90"]
            description += f" (90%+ bonus!)"
        
        # Get current XP and level
        cursor.execute("""
            SELECT total_xp, level, daily_goal_minutes, today_minutes, goal_date
            FROM user_xp WHERE user_id = %s
        """, (user_id,))
        
        row = cursor.fetchone()
        
        if row:
            old_total_xp, old_level, daily_goal_minutes, today_minutes, goal_date = row
            # Reset today_minutes if it's a new day
            if goal_date != today:
                today_minutes = 0
        else:
            old_total_xp, old_level = 0, 1
            daily_goal_minutes, today_minutes = 10, 0
        
        new_total_xp = old_total_xp + xp_earned
        new_level = calculate_level(new_total_xp)
        level_up = new_level > old_level
        
        # Update today_minutes if minutes_spent provided
        new_today_minutes = today_minutes + minutes_spent
        
        # Check if daily goal was just completed
        daily_goal_just_completed = False
        if daily_goal_minutes > 0:
            was_complete = today_minutes >= daily_goal_minutes
            is_complete = new_today_minutes >= daily_goal_minutes
            if not was_complete and is_complete:
                daily_goal_just_completed = True
                # Award daily goal bonus
                xp_earned += XP_VALUES["daily_goal_complete"]
                new_total_xp += XP_VALUES["daily_goal_complete"]
                new_level = calculate_level(new_total_xp)
                level_up = new_level > old_level
        
        # Upsert user XP
        cursor.execute("""
            INSERT INTO user_xp (user_id, total_xp, level, daily_goal_minutes, today_minutes, goal_date)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (user_id) DO UPDATE SET
                total_xp = EXCLUDED.total_xp,
                level = EXCLUDED.level,
                today_minutes = EXCLUDED.today_minutes,
                goal_date = EXCLUDED.goal_date,
                updated_at = now()
        """, (user_id, new_total_xp, new_level, daily_goal_minutes, new_today_minutes, today))
        
        # Record XP history
        cursor.execute("""
            INSERT INTO xp_history (user_id, xp_amount, activity_type, activity_id, description)
            VALUES (%s, %s, %s, %s, %s)
        """, (user_id, xp_earned, activity_type, activity_id, description))
        
        # If daily goal was completed, add a separate history entry
        if daily_goal_just_completed:
            cursor.execute("""
                INSERT INTO xp_history (user_id, xp_amount, activity_type, description)
                VALUES (%s, %s, 'daily_goal_complete', 'Daily Goal Completed!')
            """, (user_id, XP_VALUES["daily_goal_complete"]))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [XP] Awarded {xp_earned} XP to user {user_id} for {activity_type}")
        
        return {
            "xp_earned": xp_earned,
            "total_xp": new_total_xp,
            "level": new_level,
            "level_up": level_up,
            "new_level": new_level if level_up else None,
            "daily_goal_just_completed": daily_goal_just_completed
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [XP] Failed to award XP: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to award XP: {str(e)}"
        )


@app.get("/api/xp/history", tags=["XP"])
async def get_xp_history(
    limit: int = 20,
    authorization: Optional[str] = Header(None)
):
    """
    Get the user's recent XP history.
    
    Query params:
    - limit: Number of records to return (default 20, max 100)
    
    Returns list of XP events with type, amount, description, and timestamp.
    """
    try:
        user_id = await verify_user(authorization)
        
        limit = min(limit, 100)  # Cap at 100
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT xp_amount, activity_type, activity_id, description, created_at
            FROM xp_history
            WHERE user_id = %s
            ORDER BY created_at DESC
            LIMIT %s
        """, (user_id, limit))
        
        rows = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        history = []
        for row in rows:
            history.append({
                "xp_amount": row[0],
                "activity_type": row[1],
                "activity_id": row[2],
                "description": row[3],
                "created_at": row[4].isoformat() if row[4] else None
            })
        
        return {"history": history}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [XP] Failed to get XP history: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get XP history: {str(e)}"
        )


@app.patch("/api/xp/daily-goal", tags=["XP"])
async def update_daily_goal(
    request: dict,
    authorization: Optional[str] = Header(None)
):
    """
    Update the user's daily learning goal.
    
    Request body:
    - daily_goal_minutes: 0 (disabled), 5, 10, 15, or 20
    """
    try:
        user_id = await verify_user(authorization)
        
        daily_goal_minutes = request.get("daily_goal_minutes")
        
        if daily_goal_minutes not in [0, 5, 10, 15, 20]:
            raise HTTPException(status_code=400, detail="daily_goal_minutes must be 0, 5, 10, 15, or 20")
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT INTO user_xp (user_id, daily_goal_minutes)
            VALUES (%s, %s)
            ON CONFLICT (user_id) DO UPDATE SET
                daily_goal_minutes = EXCLUDED.daily_goal_minutes,
                updated_at = now()
            RETURNING daily_goal_minutes
        """, (user_id, daily_goal_minutes))
        
        result = cursor.fetchone()
        conn.commit()
        cursor.close()
        conn.close()
        
        return {"daily_goal_minutes": result[0]}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [XP] Failed to update daily goal: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to update daily goal: {str(e)}"
        )


# ==========================================
# Spaced Repetition (SM-2 Algorithm)
# ==========================================

def calculate_sm2(quality: int, easiness_factor: float, interval: int, repetition: int):
    """
    Implementation of the SM-2 spaced repetition algorithm.
    
    Args:
        quality: User's rating 0-5 (0-2 = forgot, 3 = hard, 4 = good, 5 = easy)
        easiness_factor: Current EF (starts at 2.5)
        interval: Current interval in days
        repetition: Number of successful repetitions
    
    Returns:
        tuple: (new_ef, new_interval, new_repetition)
    """
    # Calculate new easiness factor
    new_ef = easiness_factor + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02))
    new_ef = max(1.3, new_ef)  # EF should not go below 1.3
    
    if quality < 3:
        # Incorrect response - reset repetitions, start over
        new_repetition = 0
        new_interval = 1
    else:
        # Correct response
        new_repetition = repetition + 1
        if new_repetition == 1:
            new_interval = 1
        elif new_repetition == 2:
            new_interval = 6
        else:
            new_interval = int(interval * new_ef)
    
    return new_ef, new_interval, new_repetition


@app.get("/api/flashcards/due", tags=["Spaced Repetition"])
async def get_due_flashcards(
    deck_id: Optional[str] = None,
    limit: int = 20,
    authorization: Optional[str] = Header(None)
):
    """
    Get flashcards that are due for review.
    
    Query params:
    - deck_id: Optional filter by deck (e.g., "greetings", "numbers")
    - limit: Max cards to return (default 20)
    
    Returns cards where next_review <= now, or cards not yet reviewed.
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        if deck_id:
            # Get due cards for specific deck
            cursor.execute("""
                SELECT card_id, deck_id, easiness_factor, interval, repetition, 
                       last_review, next_review, total_reviews, correct_count, incorrect_count
                FROM spaced_repetition
                WHERE user_id = %s 
                AND deck_id = %s
                AND (next_review IS NULL OR next_review <= NOW())
                ORDER BY next_review ASC NULLS FIRST
                LIMIT %s
            """, (user_id, deck_id, limit))
        else:
            # Get due cards across all decks
            cursor.execute("""
                SELECT card_id, deck_id, easiness_factor, interval, repetition,
                       last_review, next_review, total_reviews, correct_count, incorrect_count
                FROM spaced_repetition
                WHERE user_id = %s 
                AND (next_review IS NULL OR next_review <= NOW())
                ORDER BY next_review ASC NULLS FIRST
                LIMIT %s
            """, (user_id, limit))
        
        rows = cursor.fetchall()
        
        # Also get count of total due cards
        if deck_id:
            cursor.execute("""
                SELECT COUNT(*) FROM spaced_repetition
                WHERE user_id = %s AND deck_id = %s
                AND (next_review IS NULL OR next_review <= NOW())
            """, (user_id, deck_id))
        else:
            cursor.execute("""
                SELECT COUNT(*) FROM spaced_repetition
                WHERE user_id = %s
                AND (next_review IS NULL OR next_review <= NOW())
            """, (user_id,))
        
        total_due = cursor.fetchone()[0]
        
        cursor.close()
        conn.close()
        
        cards = []
        for row in rows:
            cards.append({
                "card_id": row[0],
                "deck_id": row[1],
                "easiness_factor": round(float(row[2]), 2),
                "interval": row[3],
                "repetition": row[4],
                "last_review": row[5].isoformat() if row[5] else None,
                "next_review": row[6].isoformat() if row[6] else None,
                "total_reviews": row[7],
                "correct_count": row[8],
                "incorrect_count": row[9],
            })
        
        return {
            "due_cards": cards,
            "total_due": total_due,
            "has_due_cards": len(cards) > 0
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [SR] Failed to get due cards: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get due cards: {str(e)}"
        )


@app.post("/api/flashcards/review", tags=["Spaced Repetition"])
async def record_flashcard_review(
    request: dict,
    authorization: Optional[str] = Header(None)
):
    """
    Record a flashcard review and update SM-2 scheduling.
    
    Request body:
    - card_id: Unique card identifier (e.g., "greetings:3")
    - deck_id: Deck identifier (e.g., "greetings")
    - quality: User's rating 0-5 (0-2 = forgot, 3 = hard, 4 = good, 5 = easy)
    
    Returns updated card state with next review date.
    """
    try:
        user_id = await verify_user(authorization)
        
        card_id = request.get("card_id")
        deck_id = request.get("deck_id")
        quality = request.get("quality", 3)  # Default to "hard"
        
        if not card_id or not deck_id:
            raise HTTPException(status_code=400, detail="card_id and deck_id are required")
        
        if quality < 0 or quality > 5:
            raise HTTPException(status_code=400, detail="quality must be 0-5")
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Get current card state (or defaults for new card)
        cursor.execute("""
            SELECT easiness_factor, interval, repetition
            FROM spaced_repetition
            WHERE user_id = %s AND card_id = %s
        """, (user_id, card_id))
        
        row = cursor.fetchone()
        
        if row:
            current_ef, current_interval, current_rep = row
        else:
            # New card - use defaults
            current_ef, current_interval, current_rep = 2.5, 1, 0
        
        # Calculate new values using SM-2
        new_ef, new_interval, new_rep = calculate_sm2(
            quality, float(current_ef), current_interval, current_rep
        )
        
        # Calculate next review date
        from datetime import datetime, timedelta
        now = datetime.now()
        next_review = now + timedelta(days=new_interval)
        
        # Track correct/incorrect
        is_correct = quality >= 3
        correct_increment = 1 if is_correct else 0
        incorrect_increment = 0 if is_correct else 1
        
        # Upsert the card progress
        cursor.execute("""
            INSERT INTO spaced_repetition (
                user_id, card_id, deck_id, easiness_factor, interval, repetition,
                last_review, next_review, total_reviews, correct_count, incorrect_count
            )
            VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s, 1, %s, %s)
            ON CONFLICT (user_id, card_id) DO UPDATE SET
                easiness_factor = EXCLUDED.easiness_factor,
                interval = EXCLUDED.interval,
                repetition = EXCLUDED.repetition,
                last_review = NOW(),
                next_review = EXCLUDED.next_review,
                total_reviews = spaced_repetition.total_reviews + 1,
                correct_count = spaced_repetition.correct_count + %s,
                incorrect_count = spaced_repetition.incorrect_count + %s,
                updated_at = NOW()
            RETURNING total_reviews, correct_count, incorrect_count
        """, (
            user_id, card_id, deck_id, new_ef, new_interval, new_rep,
            next_review, correct_increment, incorrect_increment,
            correct_increment, incorrect_increment
        ))
        
        result = cursor.fetchone()
        conn.commit()
        cursor.close()
        conn.close()
        
        return {
            "card_id": card_id,
            "deck_id": deck_id,
            "quality": quality,
            "is_correct": is_correct,
            "easiness_factor": round(new_ef, 2),
            "interval_days": new_interval,
            "repetition": new_rep,
            "next_review": next_review.isoformat(),
            "total_reviews": result[0],
            "correct_count": result[1],
            "incorrect_count": result[2],
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [SR] Failed to record review: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to record review: {str(e)}"
        )


@app.get("/api/flashcards/stats/summary", tags=["Spaced Repetition"])
async def get_spaced_repetition_summary(
    authorization: Optional[str] = Header(None)
):
    """
    Get summary statistics for spaced repetition progress.
    
    Returns:
    - total_cards: Total cards the user has reviewed at least once
    - due_today: Cards due for review today
    - mastered: Cards with high EF (>2.3) and long interval (>30 days)
    - learning: Cards currently being learned (interval < 7 days)
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Get summary stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_cards,
                COUNT(*) FILTER (WHERE next_review IS NULL OR next_review <= NOW()) as due_today,
                COUNT(*) FILTER (WHERE easiness_factor > 2.3 AND interval > 30) as mastered,
                COUNT(*) FILTER (WHERE interval < 7) as learning
            FROM spaced_repetition
            WHERE user_id = %s
        """, (user_id,))
        
        row = cursor.fetchone()
        
        cursor.close()
        conn.close()
        
        if row:
            return {
                "total_cards": row[0],
                "due_today": row[1],
                "mastered": row[2],
                "learning": row[3],
                "has_cards": row[0] > 0
            }
        
        return {
            "total_cards": 0,
            "due_today": 0,
            "mastered": 0,
            "learning": 0,
            "has_cards": False
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [SR] Failed to get summary: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get summary: {str(e)}"
        )


@app.get("/api/subscription/status", response_model=SubscriptionStatusResponse, tags=["Usage"])
async def get_subscription_status(
    authorization: Optional[str] = Header(None)
):
    """
    Get the current user's subscription status.
    
    Returns whether the user has a premium subscription and which features are enabled.
    Checks Clerk Billing for active subscriptions.
    """
    try:
        # Verify user
        user_id = await verify_user(authorization)
        
        is_premium = False
        plan_name = None
        features = []
        
        # Check Clerk for subscription status
        if clerk:
            try:
                import httpx
                
                # Method 1: Call Clerk's billing REST API directly
                clerk_secret_key = os.getenv("CLERK_SECRET_KEY")
                if clerk_secret_key:
                    try:
                        # Call the subscriptions endpoint directly
                        headers = {
                            "Authorization": f"Bearer {clerk_secret_key}",
                            "Content-Type": "application/json"
                        }
                        
                        # Get user's subscriptions from Clerk's billing API
                        response = httpx.get(
                            f"https://api.clerk.com/v1/users/{user_id}/subscriptions",
                            headers=headers,
                            timeout=10.0
                        )
                        
                        logger.info(f"🔍 [SUBSCRIPTION] Clerk API response status: {response.status_code}")
                        
                        if response.status_code == 200:
                            subscriptions = response.json()
                            logger.info(f"🔍 [SUBSCRIPTION] Subscriptions data: {subscriptions}")
                            
                            # Check if user has any active subscriptions
                            if isinstance(subscriptions, list) and len(subscriptions) > 0:
                                for sub in subscriptions:
                                    status = sub.get('status', '')
                                    if status in ['active', 'trialing']:
                                        is_premium = True
                                        plan_name = sub.get('plan', {}).get('name', 'Premium') if isinstance(sub.get('plan'), dict) else 'Premium'
                                        features = ['unlimited_chat', 'unlimited_games', 'unlimited_quizzes']
                                        logger.info(f"✅ [SUBSCRIPTION] Found active subscription! Status: {status}, Plan: {plan_name}")
                                        break
                            elif isinstance(subscriptions, dict):
                                # Single subscription object
                                status = subscriptions.get('status', '')
                                if status in ['active', 'trialing']:
                                    is_premium = True
                                    plan_name = subscriptions.get('plan', {}).get('name', 'Premium')
                                    features = ['unlimited_chat', 'unlimited_games', 'unlimited_quizzes']
                        else:
                            logger.info(f"ℹ️ [SUBSCRIPTION] Clerk API returned {response.status_code}: {response.text[:200]}")
                            
                    except Exception as api_err:
                        logger.info(f"ℹ️ [SUBSCRIPTION] Direct API call failed: {api_err}")
                
                # Method 2: Fall back to user metadata check (in case Clerk stores it there in future)
                if not is_premium:
                    user = clerk.users.get(user_id=user_id)
                    public_metadata = getattr(user, 'public_metadata', {}) or {}
                    
                    # Check for subscription in metadata
                    subscription = public_metadata.get('subscription')
                    plan = public_metadata.get('plan')
                    
                    if subscription and isinstance(subscription, dict):
                        status = subscription.get('status', '')
                        if status in ['active', 'trialing']:
                            is_premium = True
                            plan_name = subscription.get('plan_name', 'Premium')
                            features = ['unlimited_chat', 'unlimited_games', 'unlimited_quizzes']
                
                logger.info(f"✅ [SUBSCRIPTION] User {user_id[:8]}... - Premium: {is_premium}, Plan: {plan_name}")
                
            except Exception as clerk_error:
                logger.error(f"⚠️ [SUBSCRIPTION] Failed to check Clerk: {clerk_error}")
                # Don't fail the request, just return free status
        
        return SubscriptionStatusResponse(
            is_premium=is_premium,
            plan_name=plan_name,
            features=features
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [SUBSCRIPTION] Failed to get status: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get subscription status: {str(e)}"
        )


# ==========================================
# Clerk Webhook Endpoint
# ==========================================

# Get webhook secret from environment
CLERK_WEBHOOK_SECRET = os.getenv("CLERK_WEBHOOK_SECRET", "")

@app.post("/api/webhooks/clerk", tags=["Webhooks"])
async def clerk_webhook(request: Request):
    """
    Handle Clerk webhook events for subscription changes.
    
    This endpoint is called by Clerk when subscription events occur:
    - subscription.created: User subscribes to a plan
    - subscription.updated: Subscription is modified
    - subscription.deleted: Subscription is cancelled
    
    The webhook updates user's publicMetadata with is_premium status.
    """
    try:
        # Get the raw body and headers
        body = await request.body()
        headers = dict(request.headers)
        
        logger.info(f"📨 [WEBHOOK] Received Clerk webhook")
        
        # Verify webhook signature if secret is configured
        if CLERK_WEBHOOK_SECRET:
            try:
                from svix.webhooks import Webhook
                
                svix_id = headers.get("svix-id")
                svix_timestamp = headers.get("svix-timestamp")
                svix_signature = headers.get("svix-signature")
                
                if not all([svix_id, svix_timestamp, svix_signature]):
                    logger.warning("⚠️ [WEBHOOK] Missing Svix headers")
                    raise HTTPException(status_code=400, detail="Missing webhook headers")
                
                wh = Webhook(CLERK_WEBHOOK_SECRET)
                payload = wh.verify(body, {
                    "svix-id": svix_id,
                    "svix-timestamp": svix_timestamp,
                    "svix-signature": svix_signature,
                })
                logger.info("✅ [WEBHOOK] Signature verified")
            except Exception as e:
                logger.error(f"❌ [WEBHOOK] Signature verification failed: {e}")
                raise HTTPException(status_code=401, detail="Invalid webhook signature")
        else:
            # No secret configured - parse body directly (for testing)
            payload = json.loads(body)
            logger.warning("⚠️ [WEBHOOK] No CLERK_WEBHOOK_SECRET configured - skipping signature verification")
        
        # Extract event type and data
        event_type = payload.get("type", "")
        data = payload.get("data", {})
        
        logger.info(f"📋 [WEBHOOK] Event type: {event_type}")
        logger.info(f"📋 [WEBHOOK] Data keys: {list(data.keys())}")
        
        # Handle subscription events
        if event_type.startswith("subscription."):
            # Clerk Billing: The actual Clerk user_id is in payer.user_id, NOT payer_id
            # payer_id is the commerce payer ID (cpayer_xxx), not the Clerk user ID (user_xxx)
            payer = data.get("payer", {}) or {}
            user_id = payer.get("user_id") or data.get("user_id")
            status = data.get("status", "")
            
            # Get plan name from items array (Clerk Billing structure)
            items = data.get("items", [])
            plan_name = "Premium"
            if items and len(items) > 0:
                first_item = items[0]
                plan_info = first_item.get("plan", {})
                plan_name = plan_info.get("name", "Premium") if isinstance(plan_info, dict) else "Premium"
            
            logger.info(f"🔔 [WEBHOOK] Subscription event for user {user_id}")
            logger.info(f"   Plan: {plan_name}, Status: {status}")
            logger.info(f"   Items: {items}")
            logger.info(f"   Payer: {payer}")
            
            if not user_id:
                logger.warning("⚠️ [WEBHOOK] No user_id/payer_id in subscription event")
                logger.warning(f"   Full data: {json.dumps(data, default=str)[:500]}")
                return {"received": True, "processed": False, "reason": "no_user_id"}
            
            if not clerk:
                logger.error("❌ [WEBHOOK] Clerk not initialized")
                return {"received": True, "processed": False, "reason": "clerk_not_initialized"}
            
            try:
                # Determine if user should be premium
                is_premium = event_type in ["subscription.created", "subscription.updated", "subscription.active"] and status in ["active", "trialing"]
                
                # Update user's public metadata
                if is_premium:
                    new_metadata = {
                        "is_premium": True,
                        "plan_name": plan_name,
                        "subscription_status": status,
                        "updated_at": datetime.utcnow().isoformat()
                    }
                    logger.info(f"✅ [WEBHOOK] Setting premium status for user {user_id}")
                else:
                    new_metadata = {
                        "is_premium": False,
                        "plan_name": None,
                        "subscription_status": status,
                        "updated_at": datetime.utcnow().isoformat()
                    }
                    logger.info(f"🔄 [WEBHOOK] Removing premium status for user {user_id}")
                
                # Update via Clerk API
                clerk.users.update(
                    user_id=user_id,
                    public_metadata=new_metadata
                )
                
                logger.info(f"✅ [WEBHOOK] Updated user {user_id} metadata: is_premium={is_premium}")
                
                return {
                    "received": True,
                    "processed": True,
                    "event_type": event_type,
                    "user_id": user_id,
                    "is_premium": is_premium
                }
                
            except Exception as e:
                logger.error(f"❌ [WEBHOOK] Failed to update user metadata: {e}")
                return {"received": True, "processed": False, "reason": str(e)}
        
        # Handle user.deleted event - clean up user data from our database
        elif event_type == "user.deleted":
            user_id = data.get("id")  # Clerk sends user data with "id" field
            
            if not user_id:
                logger.warning("⚠️ [WEBHOOK] No user_id in user.deleted event")
                return {"received": True, "processed": False, "reason": "no_user_id"}
            
            logger.info(f"🗑️ [WEBHOOK] User deleted event for user: {user_id}")
            
            try:
                # Delete all user data from our database
                db_url = os.getenv("DATABASE_URL")
                import psycopg2
                conn = psycopg2.connect(db_url)
                cursor = conn.cursor()
                
                deleted_counts = {}
                
                # Delete shared conversations first (foreign key to conversations)
                cursor.execute("""
                    DELETE FROM shared_conversations
                    WHERE user_id = %s
                """, (user_id,))
                deleted_counts["shared_conversations"] = cursor.rowcount
                
                # Delete conversation logs
                cursor.execute("""
                    DELETE FROM conversation_logs
                    WHERE conversation_id IN (
                        SELECT id FROM conversations WHERE user_id = %s
                    )
                """, (user_id,))
                deleted_counts["conversation_logs"] = cursor.rowcount
                
                # Delete conversations
                cursor.execute("""
                    DELETE FROM conversations
                    WHERE user_id = %s
                """, (user_id,))
                deleted_counts["conversations"] = cursor.rowcount
                
                # Delete quiz results
                cursor.execute("""
                    DELETE FROM quiz_results
                    WHERE user_id = %s
                """, (user_id,))
                deleted_counts["quiz_results"] = cursor.rowcount
                
                # Delete game results
                cursor.execute("""
                    DELETE FROM game_results
                    WHERE user_id = %s
                """, (user_id,))
                deleted_counts["game_results"] = cursor.rowcount
                
                # Delete daily usage
                cursor.execute("""
                    DELETE FROM user_daily_usage
                    WHERE user_id = %s
                """, (user_id,))
                deleted_counts["daily_usage"] = cursor.rowcount
                
                # Delete saved flashcard decks (if table exists)
                try:
                    cursor.execute("""
                        DELETE FROM saved_flashcard_decks
                        WHERE user_id = %s
                    """, (user_id,))
                    deleted_counts["flashcard_decks"] = cursor.rowcount
                except Exception:
                    pass  # Table may not exist
                
                # Delete XP data (if table exists)
                try:
                    cursor.execute("""
                        DELETE FROM user_xp
                        WHERE user_id = %s
                    """, (user_id,))
                    deleted_counts["xp"] = cursor.rowcount
                except Exception:
                    pass  # Table may not exist
                
                # Delete learning activities (if table exists)
                try:
                    cursor.execute("""
                        DELETE FROM learning_activities
                        WHERE user_id = %s
                    """, (user_id,))
                    deleted_counts["learning_activities"] = cursor.rowcount
                except Exception:
                    pass  # Table may not exist
                
                # Delete topic progress (if table exists)
                try:
                    cursor.execute("""
                        DELETE FROM topic_progress
                        WHERE user_id = %s
                    """, (user_id,))
                    deleted_counts["topic_progress"] = cursor.rowcount
                except Exception:
                    pass  # Table may not exist
                
                conn.commit()
                cursor.close()
                conn.close()
                
                logger.info(f"✅ [WEBHOOK] Deleted user data for {user_id}: {deleted_counts}")
                
                return {
                    "received": True,
                    "processed": True,
                    "event_type": event_type,
                    "user_id": user_id,
                    "deleted": deleted_counts
                }
                
            except Exception as e:
                logger.error(f"❌ [WEBHOOK] Failed to delete user data: {e}")
                return {"received": True, "processed": False, "reason": str(e)}
        
        # Handle other event types (log but don't process)
        logger.info(f"ℹ️ [WEBHOOK] Unhandled event type: {event_type}")
        return {"received": True, "processed": False, "reason": "unhandled_event_type"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [WEBHOOK] Error processing webhook: {e}")
        raise HTTPException(status_code=500, detail=f"Webhook processing failed: {str(e)}")


# ==========================================
# Vocabulary Browser Endpoints
# ==========================================

from .dictionary_service import get_dictionary_service

@app.get("/api/vocabulary/categories", tags=["Vocabulary"])
async def get_vocabulary_categories():
    """
    Get all vocabulary categories with word counts.
    
    Returns a list of categories like Greetings, Family, Numbers, etc.
    """
    try:
        service = get_dictionary_service()
        categories = service.get_categories()
        stats = service.get_stats()
        
        return {
            "categories": categories,
            "total_words": stats["total_words"]
        }
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get categories: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/categories/{category_id}", tags=["Vocabulary"])
async def get_category_words(
    category_id: str,
    limit: int = 100,
    offset: int = 0
):
    """
    Get all words in a specific category.
    
    Args:
        category_id: The category ID (e.g., 'greetings', 'family', 'numbers')
        limit: Maximum number of words to return (default 100)
        offset: Number of words to skip (for pagination)
    """
    try:
        service = get_dictionary_service()
        result = service.get_category_words(category_id, limit=limit, offset=offset)
        
        if not result["category"]:
            raise HTTPException(status_code=404, detail=f"Category '{category_id}' not found")
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get category words: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/search", tags=["Vocabulary"])
async def search_vocabulary(
    q: str,
    limit: int = 50
):
    """
    Search for words in Chamorro or English.
    
    Args:
        q: Search query (minimum 2 characters)
        limit: Maximum number of results (default 50)
    """
    try:
        if len(q) < 2:
            return {"results": [], "query": q, "total": 0}
        
        service = get_dictionary_service()
        results = service.search(q, limit=limit)
        
        return {
            "results": results,
            "query": q,
            "total": len(results)
        }
    except Exception as e:
        logger.error(f"❌ [VOCAB] Search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/word/{word}", tags=["Vocabulary"])
async def get_word_details(word: str, enhanced: bool = True):
    """
    Get detailed information for a specific word.
    
    Args:
        word: The Chamorro word to look up
        enhanced: If True (default), use morphological analysis to find root words
                  and provide helpful context even if exact word not found
    
    Returns (when enhanced=True):
        - found: bool - whether any definition was found
        - word: str - the original word
        - definition: dict or None - the dictionary entry
        - root_word: str or None - the root word if different from original
        - morphology_note: str or None - explanation of the word form
        - suggestions: list - other possible lookups to try
    """
    try:
        service = get_dictionary_service()
        
        if enhanced:
            result = service.get_word_with_morphology(word)
            return result
        else:
            result = service.get_word(word)
            if not result:
                raise HTTPException(status_code=404, detail=f"Word '{word}' not found")
            return result
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get word: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/stats", tags=["Vocabulary"])
async def get_vocabulary_stats():
    """
    Get dictionary statistics.
    """
    try:
        service = get_dictionary_service()
        return service.get_stats()
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/flashcards/{category_id}", tags=["Vocabulary"])
async def get_vocabulary_flashcards(
    category_id: str,
    count: int = 10,
    shuffle: bool = True
):
    """
    Get flashcard-formatted words from a category.
    
    Returns words formatted for flashcard display with:
    - front: Chamorro word
    - back: English definition
    - part_of_speech: Part of speech
    - example: Example sentence (if available)
    
    Args:
        category_id: The category to get words from (e.g., 'greetings', 'family')
        count: Number of flashcards to return (default 10)
        shuffle: Whether to shuffle cards for variety (default True)
    """
    try:
        service = get_dictionary_service()
        result = service.get_flashcards(category_id, count=count, shuffle=shuffle)
        
        if not result["category"]:
            raise HTTPException(status_code=404, detail=f"Category '{category_id}' not found")
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get flashcards: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/word-of-the-day", tags=["Vocabulary"])
async def get_word_of_the_day():
    """
    Get the word of the day.
    
    Returns a deterministic word based on the current date,
    so everyone sees the same word on the same day.
    
    Response includes:
    - chamorro: The Chamorro word
    - english: English definition
    - part_of_speech: Part of speech (n., v., adj., etc.)
    - example: Example sentence (if available)
    - category: Word category
    - date: Current date (ISO format)
    """
    try:
        service = get_dictionary_service()
        return service.get_word_of_the_day()
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to get word of the day: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/vocabulary/quiz/{category_id}", tags=["Vocabulary"])
async def generate_quiz_from_dictionary(
    category_id: str,
    count: int = 10,
    types: str = "multiple_choice,type_answer"
):
    """
    Generate quiz questions from dictionary words.
    
    Args:
        category_id: Category to generate questions from (e.g., 'greetings', 'family', 'all')
        count: Number of questions to generate (default 10)
        types: Comma-separated question types (multiple_choice, type_answer)
    
    Response includes:
    - questions: Array of quiz questions
    - total: Number of questions generated
    - category: Category name
    - available_words: Total words available in this category
    """
    try:
        service = get_dictionary_service()
        question_types = [t.strip() for t in types.split(",")]
        result = service.generate_quiz_questions(category_id, count=count, question_types=question_types)
        
        if not result["questions"]:
            raise HTTPException(status_code=404, detail=f"Category '{category_id}' not found or has insufficient words")
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [VOCAB] Failed to generate quiz: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# Story Mode Endpoints (Pre-extracted from Lengguahi-ta)
# =====================================================

from .story_service import get_available_stories as get_stories_list, get_story, get_story_categories

@app.get("/api/stories/available", tags=["Stories"])
async def list_available_stories(
    limit: int = 50,
    category: Optional[str] = None
):
    """
    Get list of available pre-extracted Chamorro stories.
    
    Stories are pre-extracted from Lengguahi-ta for instant loading.
    
    Args:
        limit: Maximum number of stories to return (default 50)
        category: Filter by category (story, legend, lesson, etc.)
    
    Response includes:
    - id: Story ID
    - title: English title
    - titleChamorro: Chamorro title
    - author: Story author
    - difficulty: beginner/intermediate/advanced
    - category: Story category
    - paragraphCount: Number of paragraphs
    - sourceUrl: Original source URL
    """
    try:
        stories = get_stories_list(category=category, limit=limit)
        
        # Group by category for better UX
        categories = {}
        for story in stories:
            cat = story.get("category", "story")
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(story)
        
        return {
            "stories": stories,
            "total": len(stories),
            "by_category": categories
        }
    except Exception as e:
        logger.error(f"❌ [STORIES] Failed to list stories: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/stories/{story_id}", tags=["Stories"])
async def get_story_endpoint(story_id: str):
    """
    Get a pre-extracted story by ID.
    
    Returns the complete story with all paragraphs (Chamorro + English).
    Word translations can be looked up via the /api/vocabulary/word/{word} endpoint.
    
    Args:
        story_id: The story ID (URL slug)
    
    Response includes:
    - id: Story ID
    - title: Chamorro title
    - titleEnglish: English title
    - author: Story author
    - description: Brief description
    - difficulty: beginner/intermediate/advanced
    - category: Story category
    - paragraphs: Array of {id, chamorro, english}
    - paragraphCount: Number of paragraphs
    - wordCount: Total words in Chamorro text
    - readingTime: Estimated reading time in minutes
    - sourceUrl: Original URL
    - attribution: Attribution text
    
    Note: This endpoint is instant - no AI generation required!
    """
    try:
        logger.info(f"📖 [STORIES] Fetching story: {story_id}")
        story = get_story(story_id)
        
        if not story:
            raise HTTPException(status_code=404, detail=f"Story '{story_id}' not found")
        
        logger.info(f"✅ [STORIES] Returned story: {story_id} ({story.get('paragraphCount', 0)} paragraphs)")
        return story
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [STORIES] Failed to get story: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/stories/categories", tags=["Stories"])
async def list_story_categories():
    """Get list of story categories with counts."""
    try:
        categories = get_story_categories()
        return {"categories": categories}
    except Exception as e:
        logger.error(f"❌ [STORIES] Failed to list categories: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================
# CONVERSATION PRACTICE ENDPOINTS
# =====================================================

from pydantic import BaseModel
from typing import List, Dict, Any

class ConversationPracticeRequest(BaseModel):
    scenario_id: str
    scenario_context: Dict[str, Any]
    conversation_history: List[Dict[str, str]]
    user_message: str
    turn_count: int
    user_id: Optional[str] = None

class ConversationPracticeResponse(BaseModel):
    chamorro_response: str
    english_translation: str
    feedback: Optional[Dict[str, Any]] = None
    objectives_completed: List[str] = []
    is_complete: bool = False
    final_score: Optional[int] = None

@app.post("/api/conversation-practice", response_model=ConversationPracticeResponse, tags=["Learning"])
async def conversation_practice(request: ConversationPracticeRequest):
    """
    AI-powered conversation practice endpoint.
    
    The AI plays a character in a scenario and responds to user input in Chamorro,
    providing gentle feedback on the user's Chamorro usage.
    """
    try:
        import openai
        
        client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Build the system prompt for the character
        context = request.scenario_context
        system_prompt = f"""You are playing the role of {context.get('character_name', 'a Chamorro speaker')} in a Chamorro language learning conversation.

SETTING: {context.get('setting', 'A conversation in Guam')}
YOUR ROLE: {context.get('character_role', 'A friendly local')}

IMPORTANT INSTRUCTIONS:
1. ALWAYS respond in Chamorro first, then provide English translation
2. Keep responses SHORT (1-3 sentences in Chamorro)
3. Stay in character - you are {context.get('character_name', 'the character')}
4. Be encouraging and patient - the user is learning
5. If the user makes mistakes, gently continue the conversation (don't correct harshly)
6. Use simple, common Chamorro phrases when possible
7. The conversation should feel natural and flow like a real interaction

OBJECTIVES the user should accomplish:
{chr(10).join('- ' + obj for obj in context.get('objectives', []))}

USEFUL PHRASES the user might use:
{chr(10).join('- ' + p for p in context.get('useful_phrases', [])[:5])}

Turn count: {request.turn_count}
Expected turns: ~10-14

If the conversation has naturally concluded (user has accomplished most objectives and said goodbye), 
set is_complete to true in your response.

RESPONSE FORMAT (JSON):
{{
    "chamorro_response": "Your response in Chamorro",
    "english_translation": "English translation of your response",
    "feedback": {{
        "corrections": ["List of gentle spelling/grammar suggestions for the USER's message, if any"],
        "encouragement": "Brief encouraging comment about their Chamorro"
    }},
    "objectives_completed": ["List of objectives the user has now completed"],
    "is_complete": false,
    "final_score": null
}}

When is_complete is true, set final_score to 1-5 based on:
- 5: Excellent - completed all objectives, good Chamorro usage
- 4: Good - completed most objectives, minor issues
- 3: Satisfactory - completed some objectives
- 2: Needs practice - struggled with objectives
- 1: Keep trying - minimal completion"""

        # Build conversation history for context
        messages = [{"role": "system", "content": system_prompt}]
        
        for msg in request.conversation_history:
            role = "assistant" if msg["role"] == "character" else "user"
            messages.append({"role": role, "content": msg["content"]})
        
        # Add the new user message
        messages.append({"role": "user", "content": request.user_message})
        
        # Call OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.7,
            max_tokens=500,
            response_format={"type": "json_object"}
        )
        
        # Parse the response
        response_text = response.choices[0].message.content
        response_data = json.loads(response_text)
        
        return ConversationPracticeResponse(
            chamorro_response=response_data.get("chamorro_response", ""),
            english_translation=response_data.get("english_translation", ""),
            feedback=response_data.get("feedback"),
            objectives_completed=response_data.get("objectives_completed", []),
            is_complete=response_data.get("is_complete", False),
            final_score=response_data.get("final_score")
        )
        
    except Exception as e:
        logger.error(f"❌ [CONVERSATION PRACTICE] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# ADMIN DASHBOARD ENDPOINTS
# =============================================================================

async def verify_admin(authorization: Optional[str] = Header(None)) -> str:
    """
    Verify that the request is from an admin user.
    Returns the user_id if valid, raises HTTPException if not.
    """
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")
    
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization format")
    
    token = authorization.replace("Bearer ", "")
    
    if not clerk:
        raise HTTPException(status_code=500, detail="Clerk not initialized")
    
    try:
        # Decode and verify the JWT
        jwks_client = clerk.jwks.get_jwks()
        keys = jwks_client.keys if hasattr(jwks_client, 'keys') else []
        
        if not keys:
            raise HTTPException(status_code=500, detail="Unable to fetch JWKS keys")
        
        from jose import jwt
        
        # Get the signing key
        unverified_header = jwt.get_unverified_header(token)
        signing_key = None
        for key in keys:
            if key.kid == unverified_header.get("kid"):
                signing_key = key
                break
        
        if not signing_key:
            raise HTTPException(status_code=401, detail="Unable to find signing key")
        
        # Verify and decode the token
        payload = jwt.decode(
            token,
            signing_key.model_dump() if hasattr(signing_key, 'model_dump') else dict(signing_key),
            algorithms=["RS256"],
            options={"verify_aud": False}
        )
        
        user_id = payload.get("sub")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token: no user ID")
        
        # Check if user is admin
        user = clerk.users.get(user_id=user_id)
        public_metadata = getattr(user, 'public_metadata', {}) or {}
        role = public_metadata.get('role')
        
        if role != 'admin':
            logger.warning(f"🚫 [ADMIN] Non-admin user {user_id} attempted admin access")
            raise HTTPException(status_code=403, detail="Admin access required")
        
        logger.info(f"✅ [ADMIN] Admin user {user_id} authenticated")
        return user_id
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] Auth error: {e}")
        raise HTTPException(status_code=401, detail=f"Authentication failed: {str(e)}")


@app.get("/api/admin/stats", response_model=AdminStatsResponse, tags=["Admin"])
async def get_admin_stats(authorization: Optional[str] = Header(None)):
    """
    Get dashboard statistics for admin panel.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        if not clerk:
            raise HTTPException(status_code=500, detail="Clerk not initialized")
        
        # Get all users from Clerk
        users_response = clerk.users.list()
        users = users_response.data if hasattr(users_response, 'data') else (users_response if isinstance(users_response, list) else [])
        
        total_users = len(users)
        premium_users = 0
        whitelisted_users = 0
        
        for user in users:
            metadata = getattr(user, 'public_metadata', {}) or {}
            if metadata.get('is_premium'):
                premium_users += 1
            if metadata.get('whitelisted') or metadata.get('is_whitelisted'):
                whitelisted_users += 1
        
        free_users = total_users - premium_users
        
        # Get database stats
        db_url = os.getenv("DATABASE_URL")
        if not db_url:
            raise HTTPException(status_code=500, detail="Database not configured")
        
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Total conversations
        cursor.execute("SELECT COUNT(*) FROM conversations WHERE deleted_at IS NULL")
        total_conversations = cursor.fetchone()[0]
        
        # Total messages
        cursor.execute("SELECT COUNT(*) FROM conversation_logs")
        total_messages = cursor.fetchone()[0]
        
        # Total quiz attempts
        cursor.execute("SELECT COUNT(*) FROM quiz_results")
        total_quiz_attempts = cursor.fetchone()[0]
        
        # Total game plays
        cursor.execute("SELECT COUNT(*) FROM game_results")
        total_game_plays = cursor.fetchone()[0]
        
        # Active today (users with activity in user_daily_usage)
        today = get_guam_date()
        cursor.execute(
            "SELECT COUNT(DISTINCT user_id) FROM user_daily_usage WHERE usage_date = %s",
            (today,)
        )
        active_today = cursor.fetchone()[0]
        
        conn.close()
        
        return AdminStatsResponse(
            total_users=total_users,
            premium_users=premium_users,
            free_users=free_users,
            whitelisted_users=whitelisted_users,
            active_today=active_today,
            total_conversations=total_conversations,
            total_messages=total_messages,
            total_quiz_attempts=total_quiz_attempts,
            total_game_plays=total_game_plays
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] Stats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/users", response_model=AdminUsersResponse, tags=["Admin"])
async def get_admin_users(
    authorization: Optional[str] = Header(None),
    page: int = Query(1, ge=1, description="Page number"),
    per_page: int = Query(20, ge=1, le=100, description="Items per page"),
    search: Optional[str] = Query(None, description="Search by email or name")
):
    """
    Get paginated list of users for admin panel.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        if not clerk:
            raise HTTPException(status_code=500, detail="Clerk not initialized")
        
        # Get users from Clerk (with pagination)
        # Note: Clerk's list has its own pagination, we'll handle it here
        offset = (page - 1) * per_page
        
        if search:
            # Search by email or name - Clerk SDK uses email_address or query param
            users_response = clerk.users.list(query=search)
        else:
            users_response = clerk.users.list()
        
        all_users = users_response.data if hasattr(users_response, 'data') else (users_response if isinstance(users_response, list) else [])
        total = len(all_users)
        
        # Apply pagination
        paginated_users = all_users[offset:offset + per_page]
        
        # Get database stats for each user
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        admin_users = []
        for user in paginated_users:
            user_id = user.id
            metadata = getattr(user, 'public_metadata', {}) or {}
            
            # Get user's email
            email = None
            if hasattr(user, 'email_addresses') and user.email_addresses:
                email = user.email_addresses[0].email_address
            
            # Get stats from database
            cursor.execute(
                "SELECT COUNT(*) FROM conversations WHERE user_id = %s AND deleted_at IS NULL",
                (user_id,)
            )
            total_conversations = cursor.fetchone()[0]
            
            cursor.execute(
                "SELECT COUNT(*) FROM conversation_logs WHERE user_id = %s",
                (user_id,)
            )
            total_messages = cursor.fetchone()[0]
            
            cursor.execute(
                "SELECT COUNT(*) FROM quiz_results WHERE user_id = %s",
                (user_id,)
            )
            total_quizzes = cursor.fetchone()[0]
            
            cursor.execute(
                "SELECT COUNT(*) FROM game_results WHERE user_id = %s",
                (user_id,)
            )
            total_games = cursor.fetchone()[0]
            
            # Get last activity timestamp (most recent of messages, quizzes, games)
            # Note: conversation_logs uses 'timestamp', others use 'created_at'
            cursor.execute("""
                SELECT MAX(last_activity) FROM (
                    SELECT MAX(timestamp) as last_activity FROM conversation_logs WHERE user_id = %s
                    UNION ALL
                    SELECT MAX(created_at) as last_activity FROM quiz_results WHERE user_id = %s
                    UNION ALL
                    SELECT MAX(created_at) as last_activity FROM game_results WHERE user_id = %s
                ) as activities
            """, (user_id, user_id, user_id))
            last_activity_row = cursor.fetchone()
            last_activity = last_activity_row[0].isoformat() if last_activity_row and last_activity_row[0] else None
            
            admin_users.append(AdminUserInfo(
                user_id=user_id,
                email=email,
                first_name=getattr(user, 'first_name', None),
                last_name=getattr(user, 'last_name', None),
                image_url=getattr(user, 'image_url', None),
                is_premium=metadata.get('is_premium', False),
                is_whitelisted=metadata.get('whitelisted', False) or metadata.get('is_whitelisted', False),
                is_banned=metadata.get('is_banned', False),
                role=metadata.get('role'),
                plan_name=metadata.get('plan_name'),
                subscription_status=metadata.get('subscription_status'),
                created_at=str(user.created_at) if hasattr(user, 'created_at') else None,
                last_sign_in=str(user.last_sign_in_at) if hasattr(user, 'last_sign_in_at') else None,
                last_activity=last_activity,
                total_conversations=total_conversations,
                total_messages=total_messages,
                total_quizzes=total_quizzes,
                total_games=total_games
            ))
        
        conn.close()
        
        total_pages = (total + per_page - 1) // per_page
        
        return AdminUsersResponse(
            users=admin_users,
            total=total,
            page=page,
            per_page=per_page,
            total_pages=total_pages
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] Users list error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/users/{user_id}", response_model=AdminUserInfo, tags=["Admin"])
async def get_admin_user(
    user_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Get detailed info for a specific user.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        if not clerk:
            raise HTTPException(status_code=500, detail="Clerk not initialized")
        
        # Get user from Clerk
        user = clerk.users.get(user_id=user_id)
        metadata = getattr(user, 'public_metadata', {}) or {}
        unsafe_metadata = getattr(user, 'unsafe_metadata', {}) or {}
        
        # Get user's email
        email = None
        if hasattr(user, 'email_addresses') and user.email_addresses:
            email = user.email_addresses[0].email_address
        
        # Get stats from database
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Total stats
        cursor.execute(
            "SELECT COUNT(*) FROM conversations WHERE user_id = %s AND deleted_at IS NULL",
            (user_id,)
        )
        total_conversations = cursor.fetchone()[0]
        
        cursor.execute(
            "SELECT COUNT(*) FROM conversation_logs WHERE user_id = %s",
            (user_id,)
        )
        total_messages = cursor.fetchone()[0]
        
        cursor.execute(
            "SELECT COUNT(*) FROM quiz_results WHERE user_id = %s",
            (user_id,)
        )
        total_quizzes = cursor.fetchone()[0]
        
        cursor.execute(
            "SELECT COUNT(*) FROM game_results WHERE user_id = %s",
            (user_id,)
        )
        total_games = cursor.fetchone()[0]
        
        # Today's usage
        today = get_guam_date()
        cursor.execute(
            "SELECT chat_count, game_count, quiz_count FROM user_daily_usage WHERE user_id = %s AND usage_date = %s",
            (user_id, today)
        )
        today_row = cursor.fetchone()
        today_chat = today_row[0] if today_row else 0
        today_games = today_row[1] if today_row else 0
        today_quizzes = today_row[2] if today_row else 0
        
        # Get last activity timestamp (most recent of messages, quizzes, games)
        cursor.execute("""
            SELECT MAX(last_activity) FROM (
                SELECT MAX(timestamp) as last_activity FROM conversation_logs WHERE user_id = %s
                UNION ALL
                SELECT MAX(created_at) as last_activity FROM quiz_results WHERE user_id = %s
                UNION ALL
                SELECT MAX(created_at) as last_activity FROM game_results WHERE user_id = %s
            ) as activities
        """, (user_id, user_id, user_id))
        last_activity_row = cursor.fetchone()
        last_activity = last_activity_row[0].isoformat() if last_activity_row and last_activity_row[0] else None
        
        conn.close()
        
        return AdminUserInfo(
            user_id=user_id,
            email=email,
            first_name=getattr(user, 'first_name', None),
            last_name=getattr(user, 'last_name', None),
            image_url=getattr(user, 'image_url', None),
            is_premium=metadata.get('is_premium', False),
            is_whitelisted=metadata.get('whitelisted', False) or metadata.get('is_whitelisted', False),
            is_banned=metadata.get('is_banned', False),
            role=metadata.get('role'),
            plan_name=metadata.get('plan_name'),
            subscription_status=metadata.get('subscription_status'),
            created_at=str(user.created_at) if hasattr(user, 'created_at') else None,
            last_sign_in=str(user.last_sign_in_at) if hasattr(user, 'last_sign_in_at') else None,
            last_activity=last_activity,
            total_conversations=total_conversations,
            total_messages=total_messages,
            total_quizzes=total_quizzes,
            total_games=total_games,
            today_chat=today_chat,
            today_games=today_games,
            today_quizzes=today_quizzes,
            # Learning preferences from unsafeMetadata
            skill_level=unsafe_metadata.get('skill_level'),
            learning_goal=unsafe_metadata.get('learning_goal'),
            onboarding_completed=unsafe_metadata.get('onboarding_completed', False)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] Get user error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/api/admin/users/{user_id}", response_model=AdminUserUpdateResponse, tags=["Admin"])
async def update_admin_user(
    user_id: str,
    request: AdminUserUpdateRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Update a user's status (premium, whitelist, role).
    Requires admin role.
    """
    admin_user_id = await verify_admin(authorization)
    
    try:
        if not clerk:
            raise HTTPException(status_code=500, detail="Clerk not initialized")
        
        # Get current user data
        user = clerk.users.get(user_id=user_id)
        current_metadata = getattr(user, 'public_metadata', {}) or {}
        
        # Build new metadata
        new_metadata = dict(current_metadata)
        
        if request.is_premium is not None:
            new_metadata['is_premium'] = request.is_premium
        
        if request.is_whitelisted is not None:
            new_metadata['is_whitelisted'] = request.is_whitelisted
            new_metadata['whitelisted'] = request.is_whitelisted  # Also set alias
            if request.is_whitelisted and not new_metadata.get('is_premium'):
                # Whitelisting automatically grants premium
                new_metadata['is_premium'] = True
        
        if request.role is not None:
            new_metadata['role'] = request.role
        
        if request.is_banned is not None:
            new_metadata['is_banned'] = request.is_banned
            if request.is_banned:
                new_metadata['banned_at'] = datetime.utcnow().isoformat()
                new_metadata['banned_by'] = admin_user_id
        
        if request.plan_name is not None:
            new_metadata['plan_name'] = request.plan_name
        elif request.is_whitelisted:
            new_metadata['plan_name'] = 'Friends & Family'
        
        new_metadata['updated_at'] = datetime.utcnow().isoformat()
        new_metadata['updated_by'] = admin_user_id
        
        # Update user in Clerk
        clerk.users.update(
            user_id=user_id,
            public_metadata=new_metadata
        )
        
        logger.info(f"✅ [ADMIN] User {user_id} updated by admin {admin_user_id}")
        logger.info(f"   New metadata: {new_metadata}")
        
        # Get updated user info
        updated_user = await get_admin_user(user_id, authorization)
        
        return AdminUserUpdateResponse(
            success=True,
            user=updated_user,
            message=f"User {user_id} updated successfully"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] Update user error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/admin/users/{user_id}/reset-onboarding", tags=["Admin"])
async def reset_user_onboarding(
    user_id: str,
    authorization: Optional[str] = Header(None)
):
    """
    Reset a user's onboarding preferences (unsafeMetadata).
    
    This clears:
    - skill_level
    - learning_goal  
    - onboarding_completed
    
    The user will see the onboarding modal on next login.
    """
    admin_user_id = await verify_admin(authorization)
    
    try:
        # Clear unsafe metadata (user preferences)
        clerk.users.update(
            user_id=user_id,
            unsafe_metadata={}
        )
        
        logger.info(f"✅ [ADMIN] Reset onboarding for user {user_id} by admin {admin_user_id}")
        
        return {
            "success": True,
            "message": f"Onboarding reset for user {user_id}. They will see onboarding on next login."
        }
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Reset onboarding error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/api/admin/users/{user_id}/preferences", tags=["Admin"])
async def update_user_preferences(
    user_id: str,
    request: dict,
    authorization: Optional[str] = Header(None)
):
    """
    Update a user's learning preferences (unsafeMetadata).
    
    Body:
    {
        "skill_level": "beginner" | "intermediate" | "advanced",
        "learning_goal": "conversation" | "culture" | "family" | "travel" | "all",
        "onboarding_completed": true | false
    }
    """
    admin_user_id = await verify_admin(authorization)
    
    try:
        # Get current unsafe metadata
        user = clerk.users.get(user_id=user_id)
        current_metadata = dict(user.unsafe_metadata) if user.unsafe_metadata else {}
        
        # Update with new values
        if 'skill_level' in request:
            current_metadata['skill_level'] = request['skill_level']
        if 'learning_goal' in request:
            current_metadata['learning_goal'] = request['learning_goal']
        if 'onboarding_completed' in request:
            current_metadata['onboarding_completed'] = request['onboarding_completed']
        
        # Update in Clerk
        clerk.users.update(
            user_id=user_id,
            unsafe_metadata=current_metadata
        )
        
        logger.info(f"✅ [ADMIN] Updated preferences for user {user_id} by admin {admin_user_id}")
        logger.info(f"   New preferences: {current_metadata}")
        
        return {
            "success": True,
            "preferences": current_metadata,
            "message": f"Preferences updated for user {user_id}"
        }
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Update preferences error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Analytics Endpoints ---

@app.get("/api/admin/analytics/usage", tags=["Admin"])
async def get_usage_trends(
    authorization: Optional[str] = Header(None),
    period: str = Query("30d", description="Period: 7d, 30d, 90d")
):
    """
    Get usage trends over time (chat, games, quizzes).
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        # Parse period
        days = {"7d": 7, "30d": 30, "90d": 90}.get(period, 30)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Get daily usage from user_daily_usage table
        cursor.execute("""
            SELECT 
                usage_date,
                SUM(chat_count) as chat_total,
                SUM(game_count) as game_total,
                SUM(quiz_count) as quiz_total,
                COUNT(DISTINCT user_id) as active_users
            FROM user_daily_usage
            WHERE usage_date >= CURRENT_DATE - INTERVAL '%s days'
            GROUP BY usage_date
            ORDER BY usage_date ASC
        """, (days,))
        
        rows = cursor.fetchall()
        
        data = []
        totals = {"chat": 0, "games": 0, "quizzes": 0, "active_users": set()}
        
        for row in rows:
            data.append({
                "date": row[0].isoformat(),
                "chat_count": row[1] or 0,
                "game_count": row[2] or 0,
                "quiz_count": row[3] or 0,
                "active_users": row[4] or 0
            })
            totals["chat"] += row[1] or 0
            totals["games"] += row[2] or 0
            totals["quizzes"] += row[3] or 0
        
        conn.close()
        
        return {
            "period": period,
            "data": data,
            "totals": {
                "chat": totals["chat"],
                "games": totals["games"],
                "quizzes": totals["quizzes"]
            }
        }
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Usage trends error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/analytics/growth", tags=["Admin"])
async def get_user_growth(
    authorization: Optional[str] = Header(None),
    period: str = Query("30d", description="Period: 7d, 30d, 90d")
):
    """
    Get user growth over time.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        days = {"7d": 7, "30d": 30, "90d": 90}.get(period, 30)
        
        if not clerk:
            raise HTTPException(status_code=500, detail="Clerk not initialized")
        
        # Get all users from Clerk
        users_response = clerk.users.list()
        users = users_response.data if hasattr(users_response, 'data') else (users_response if isinstance(users_response, list) else [])
        
        # Group users by signup date
        from collections import defaultdict
        daily_signups = defaultdict(int)
        daily_premium = defaultdict(int)
        
        today = datetime.now().date()
        cutoff = today - timedelta(days=days)
        
        for user in users:
            if hasattr(user, 'created_at') and user.created_at:
                # Clerk timestamps are in milliseconds
                signup_date = datetime.fromtimestamp(user.created_at / 1000).date()
                if signup_date >= cutoff:
                    daily_signups[signup_date.isoformat()] += 1
                    
                metadata = getattr(user, 'public_metadata', {}) or {}
                if metadata.get('is_premium'):
                    daily_premium[signup_date.isoformat()] += 1
        
        # Build cumulative data
        data = []
        total_users = len(users)
        
        # Start from cutoff date
        current_date = cutoff
        cumulative = total_users - sum(daily_signups.values())  # Users before cutoff
        
        while current_date <= today:
            date_str = current_date.isoformat()
            new_users = daily_signups.get(date_str, 0)
            cumulative += new_users
            
            data.append({
                "date": date_str,
                "total_users": cumulative,
                "new_users": new_users,
                "premium_users": daily_premium.get(date_str, 0)
            })
            current_date += timedelta(days=1)
        
        return {
            "period": period,
            "data": data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [ADMIN] User growth error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/analytics/features", tags=["Admin"])
async def get_feature_usage(
    authorization: Optional[str] = Header(None)
):
    """
    Get feature popularity breakdown.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Total counts
        cursor.execute("SELECT COUNT(*) FROM conversation_logs")
        chat_total = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM game_results")
        games_total = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM quiz_results")
        quizzes_total = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM conversations WHERE deleted_at IS NULL")
        conversations_total = cursor.fetchone()[0]
        
        # Game breakdown by type
        cursor.execute("""
            SELECT game_type, COUNT(*) as count
            FROM game_results
            GROUP BY game_type
            ORDER BY count DESC
        """)
        game_breakdown = {row[0]: row[1] for row in cursor.fetchall()}
        
        # Quiz breakdown by category
        cursor.execute("""
            SELECT category_id, COUNT(*) as count
            FROM quiz_results
            GROUP BY category_id
            ORDER BY count DESC
        """)
        quiz_breakdown = {row[0] or 'Unknown': row[1] for row in cursor.fetchall()}
        
        conn.close()
        
        return {
            "chat_total": chat_total,
            "games_total": games_total,
            "quizzes_total": quizzes_total,
            "conversations_total": conversations_total,
            "game_breakdown": game_breakdown,
            "quiz_breakdown": quiz_breakdown
        }
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Feature usage error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==============================================================================
# Admin Settings Endpoints
# ==============================================================================

@app.get("/api/admin/settings", tags=["Admin"])
async def get_admin_settings(authorization: Optional[str] = Header(None)):
    """
    Get all site settings for admin panel.
    Requires admin role.
    """
    await verify_admin(authorization)
    
    try:
        db_url = os.getenv("DATABASE_URL")
        if not db_url:
            raise HTTPException(status_code=500, detail="Database not configured")
        
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT key, value, description, updated_at, updated_by
            FROM site_settings
            ORDER BY key
        """)
        
        settings = {}
        for row in cursor.fetchall():
            settings[row[0]] = {
                "value": row[1],
                "description": row[2],
                "updated_at": row[3].isoformat() if row[3] else None,
                "updated_by": row[4]
            }
        
        cursor.close()
        conn.close()
        
        logger.info(f"✅ [ADMIN] Retrieved {len(settings)} settings")
        return {"settings": settings}
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Get settings error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/api/admin/settings", tags=["Admin"])
async def update_admin_settings(
    request: Request,
    authorization: Optional[str] = Header(None)
):
    """
    Update site settings.
    Requires admin role.
    
    Body should be a dict of {key: value} pairs to update.
    Example: {"promo_enabled": "true", "promo_end_date": "2026-01-15"}
    """
    user_id = await verify_admin(authorization)
    
    try:
        body = await request.json()
        
        if not body or not isinstance(body, dict):
            raise HTTPException(status_code=400, detail="Request body must be a JSON object")
        
        db_url = os.getenv("DATABASE_URL")
        if not db_url:
            raise HTTPException(status_code=500, detail="Database not configured")
        
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        updated_keys = []
        for key, value in body.items():
            cursor.execute("""
                UPDATE site_settings 
                SET value = %s, updated_at = NOW(), updated_by = %s
                WHERE key = %s
            """, (str(value), user_id, key))
            
            if cursor.rowcount > 0:
                updated_keys.append(key)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        # Clear the settings cache so changes take effect immediately
        _clear_settings_cache()
        
        logger.info(f"✅ [ADMIN] Updated settings: {updated_keys} by {user_id}")
        return {"success": True, "updated": updated_keys}
        
    except Exception as e:
        logger.error(f"❌ [ADMIN] Update settings error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==============================================================================
# Site Settings Helper Functions
# ==============================================================================

# Cache for site settings (to avoid DB queries on every request)
_settings_cache: dict = {}
_settings_cache_time: float = 0
_SETTINGS_CACHE_TTL = 60  # Cache for 60 seconds

def _clear_settings_cache():
    """Clear the settings cache (called when settings are updated)."""
    global _settings_cache, _settings_cache_time
    _settings_cache = {}
    _settings_cache_time = 0

def get_site_setting(key: str, default: str = None) -> str:
    """
    Get a site setting from the database (with caching).
    Falls back to default if not found.
    """
    global _settings_cache, _settings_cache_time
    
    # Check if cache is still valid
    if time.time() - _settings_cache_time < _SETTINGS_CACHE_TTL and _settings_cache:
        return _settings_cache.get(key, default)
    
    try:
        db_url = os.getenv("DATABASE_URL")
        if not db_url:
            return default
        
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        cursor.execute("SELECT key, value FROM site_settings")
        
        _settings_cache = {row[0]: row[1] for row in cursor.fetchall()}
        _settings_cache_time = time.time()
        
        cursor.close()
        conn.close()
        
        return _settings_cache.get(key, default)
        
    except Exception as e:
        logger.warning(f"⚠️ Failed to get site setting {key}: {e}")
        return default


def is_promo_active_from_db() -> bool:
    """Check if promo is active based on database settings."""
    promo_enabled = get_site_setting("promo_enabled", "false")
    if promo_enabled.lower() != "true":
        return False
    
    promo_end_date = get_site_setting("promo_end_date", "2026-01-06")
    try:
        end_date = datetime.strptime(promo_end_date, "%Y-%m-%d").date()
        today = get_guam_date()
        return today <= end_date
    except ValueError:
        return False


# ==============================================================================
# Learning Path Endpoints
# ==============================================================================

# Define learning paths (same as frontend, keep in sync)
BEGINNER_PATH = [
    {"id": "greetings", "title": "Greetings & Basics", "description": "Learn 'Håfa Adai' and how to introduce yourself", "icon": "👋", "estimated_minutes": 5, "flashcard_category": "greetings", "quiz_category": "greetings", "level": "beginner"},
    {"id": "numbers", "title": "Numbers (1-10)", "description": "Learn to count in Chamorro", "icon": "🔢", "estimated_minutes": 5, "flashcard_category": "numbers", "quiz_category": "numbers", "level": "beginner"},
    {"id": "colors", "title": "Colors", "description": "Learn the colors of the rainbow", "icon": "🎨", "estimated_minutes": 5, "flashcard_category": "colors", "quiz_category": "colors", "level": "beginner"},
    {"id": "family", "title": "Family Members", "description": "Words for mother, father, siblings, and more", "icon": "👨‍👩‍👧‍👦", "estimated_minutes": 5, "flashcard_category": "family", "quiz_category": "family", "level": "beginner"},
    {"id": "food", "title": "Food & Drinks", "description": "Common foods and island favorites", "icon": "🍚", "estimated_minutes": 5, "flashcard_category": "food", "quiz_category": "food", "level": "beginner"},
    {"id": "animals", "title": "Animals", "description": "Learn about island creatures and pets", "icon": "🐠", "estimated_minutes": 5, "flashcard_category": "animals", "quiz_category": "animals", "level": "beginner"},
    {"id": "phrases", "title": "Common Phrases", "description": "Everyday expressions and useful phrases", "icon": "💬", "estimated_minutes": 5, "flashcard_category": "phrases", "quiz_category": "common-phrases", "level": "beginner"},
]

INTERMEDIATE_PATH = [
    {"id": "questions", "title": "Question Words", "description": "Learn to ask who, what, where, when, why, how", "icon": "❓", "estimated_minutes": 6, "flashcard_category": "questions", "quiz_category": "questions", "level": "intermediate"},
    {"id": "body-parts", "title": "Body Parts", "description": "Learn words for parts of the body", "icon": "🫀", "estimated_minutes": 6, "flashcard_category": "body", "quiz_category": "body-parts", "level": "intermediate"},
    {"id": "days", "title": "Days of the Week", "description": "Learn to say the days in Chamorro", "icon": "📅", "estimated_minutes": 5, "flashcard_category": "days", "quiz_category": "days", "level": "intermediate"},
    {"id": "months", "title": "Months & Seasons", "description": "Learn the months of the year", "icon": "🗓️", "estimated_minutes": 6, "flashcard_category": "months", "quiz_category": "months", "level": "intermediate"},
    {"id": "verbs", "title": "Common Verbs", "description": "Action words for everyday activities", "icon": "🏃", "estimated_minutes": 7, "flashcard_category": "verbs", "quiz_category": "verbs", "level": "intermediate"},
    {"id": "adjectives", "title": "Describing Things", "description": "Adjectives for size, quality, and feelings", "icon": "✨", "estimated_minutes": 6, "flashcard_category": "adjectives", "quiz_category": "adjectives", "level": "intermediate"},
    {"id": "sentences", "title": "Simple Sentences", "description": "Put words together to make sentences", "icon": "📝", "estimated_minutes": 8, "flashcard_category": "sentences", "quiz_category": "sentences", "level": "intermediate"},
]

ADVANCED_PATH = [
    {"id": "places", "title": "Places & Locations", "description": "Buildings, landmarks, and location phrases", "icon": "🏠", "estimated_minutes": 8, "flashcard_category": "places", "quiz_category": "places", "level": "advanced"},
    {"id": "weather", "title": "Weather & Nature", "description": "Weather conditions and natural environment", "icon": "🌞", "estimated_minutes": 8, "flashcard_category": "weather", "quiz_category": "weather", "level": "advanced"},
    {"id": "household", "title": "Home & Household", "description": "Rooms, furniture, and household items", "icon": "🛋️", "estimated_minutes": 8, "flashcard_category": "household", "quiz_category": "household", "level": "advanced"},
    {"id": "directions", "title": "Directions & Travel", "description": "Directions, movement, and transportation", "icon": "🧭", "estimated_minutes": 8, "flashcard_category": "directions", "quiz_category": "directions", "level": "advanced"},
    {"id": "shopping", "title": "Shopping & Money", "description": "Buying, selling, and money vocabulary", "icon": "💰", "estimated_minutes": 8, "flashcard_category": "shopping", "quiz_category": "shopping", "level": "advanced"},
    {"id": "daily-life", "title": "Work & Daily Life", "description": "Jobs, school, and daily activities", "icon": "💼", "estimated_minutes": 8, "flashcard_category": "daily-life", "quiz_category": "daily-life", "level": "advanced"},
    {"id": "culture", "title": "Culture & Celebrations", "description": "Traditions, fiestas, and respect language", "icon": "🎉", "estimated_minutes": 10, "flashcard_category": "culture", "quiz_category": "culture", "level": "advanced"},
]

# Combined all topics for lookups
ALL_TOPICS = BEGINNER_PATH + INTERMEDIATE_PATH + ADVANCED_PATH


@app.get("/api/learning/recommended", tags=["Learning Path"])
async def get_recommended_topic(
    authorization: Optional[str] = Header(None)
):
    """
    Get the recommended next topic for the user based on their progress.
    
    Returns:
    - recommendation_type: 'start', 'continue', 'next', 'review', or 'complete'
    - topic: The recommended topic (null if all complete)
    - progress: User's progress on that topic
    - completed_topics: Number of topics completed
    - message: User-friendly message
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        # Get user's progress on all topics
        cursor.execute("""
            SELECT topic_id, started_at, completed_at, best_quiz_score, 
                   flashcards_viewed, last_activity_at
            FROM user_topic_progress
            WHERE user_id = %s
        """, (user_id,))
        
        progress_rows = cursor.fetchall()
        cursor.close()
        conn.close()
        
        # Build progress map
        progress_map = {}
        for row in progress_rows:
            progress_map[row[0]] = {
                "topic_id": row[0],
                "started_at": row[1].isoformat() if row[1] else None,
                "completed_at": row[2].isoformat() if row[2] else None,
                "best_quiz_score": row[3],
                "flashcards_viewed": row[4] or 0,
                "last_activity_at": row[5].isoformat() if row[5] else None,
            }
        
        # Count completions per level
        beginner_completed = sum(
            1 for t in BEGINNER_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        intermediate_completed = sum(
            1 for t in INTERMEDIATE_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        advanced_completed = sum(
            1 for t in ADVANCED_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        total_completed = beginner_completed + intermediate_completed + advanced_completed
        beginner_complete = beginner_completed == len(BEGINNER_PATH)
        intermediate_complete = intermediate_completed == len(INTERMEDIATE_PATH)
        
        # Find recommendation
        recommendation_type = "start"
        recommended_topic = None
        topic_progress = None
        message = ""
        
        # First, check beginner path
        for topic in BEGINNER_PATH:
            topic_id = topic["id"]
            progress = progress_map.get(topic_id)
            
            if not progress:
                # Never started this topic - recommend it
                recommendation_type = "start" if total_completed == 0 else "next"
                recommended_topic = topic
                topic_progress = {
                    "topic_id": topic_id,
                    "started_at": None,
                    "completed_at": None,
                    "best_quiz_score": None,
                    "flashcards_viewed": 0,
                    "last_activity_at": None,
                }
                if total_completed == 0:
                    message = f"Start your Chamorro journey with {topic['title']}!"
                else:
                    message = f"Great progress! Ready for {topic['title']}?"
                break
            
            elif not progress["completed_at"]:
                # Started but not completed - continue
                recommendation_type = "continue"
                recommended_topic = topic
                topic_progress = progress
                message = f"Continue learning {topic['title']}"
                break
        
        # If beginner complete, check intermediate path
        if beginner_complete and not recommended_topic:
            for topic in INTERMEDIATE_PATH:
                topic_id = topic["id"]
                progress = progress_map.get(topic_id)
                
                if not progress:
                    # Never started this topic - recommend it
                    recommendation_type = "next"
                    recommended_topic = topic
                    topic_progress = {
                        "topic_id": topic_id,
                        "started_at": None,
                        "completed_at": None,
                        "best_quiz_score": None,
                        "flashcards_viewed": 0,
                        "last_activity_at": None,
                    }
                    message = f"Level up! Start {topic['title']} in Intermediate."
                    break
                
                elif not progress["completed_at"]:
                    # Started but not completed - continue
                    recommendation_type = "continue"
                    recommended_topic = topic
                    topic_progress = progress
                    message = f"Continue learning {topic['title']}"
                    break
        
        # If intermediate complete, check advanced path
        if beginner_complete and intermediate_complete and not recommended_topic:
            for topic in ADVANCED_PATH:
                topic_id = topic["id"]
                progress = progress_map.get(topic_id)
                
                if not progress:
                    # Never started this topic - recommend it
                    recommendation_type = "next"
                    recommended_topic = topic
                    topic_progress = {
                        "topic_id": topic_id,
                        "started_at": None,
                        "completed_at": None,
                        "best_quiz_score": None,
                        "flashcards_viewed": 0,
                        "last_activity_at": None,
                    }
                    message = f"You're Advanced now! Start {topic['title']}."
                    break
                
                elif not progress["completed_at"]:
                    # Started but not completed - continue
                    recommendation_type = "continue"
                    recommended_topic = topic
                    topic_progress = progress
                    message = f"Continue learning {topic['title']}"
                    break
        
        # All completed
        if not recommended_topic:
            recommendation_type = "complete"
            # Suggest reviewing the one with lowest score across all topics
            worst_topic = None
            worst_score = 101
            for topic in ALL_TOPICS:
                progress = progress_map.get(topic["id"])
                if progress and progress["best_quiz_score"] is not None:
                    if progress["best_quiz_score"] < worst_score:
                        worst_score = progress["best_quiz_score"]
                        worst_topic = topic
                        topic_progress = progress
            
            if worst_topic and worst_score < 100:
                recommendation_type = "review"
                recommended_topic = worst_topic
                message = f"You've completed all topics! Review {worst_topic['title']} to improve your score."
            else:
                message = "Congratulations! You've mastered all available topics! 🎉"
        
        return {
            "recommendation_type": recommendation_type,
            "topic": recommended_topic,
            "progress": topic_progress,
            "total_topics": len(ALL_TOPICS),
            "completed_topics": total_completed,
            "beginner_complete": beginner_complete,
            "message": message,
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Get recommended topic error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/learning/progress/{topic_id}", tags=["Learning Path"])
async def update_topic_progress(
    topic_id: str,
    action: str = Query(..., description="Action: 'start', 'flashcard_viewed', 'quiz_completed'"),
    quiz_score: Optional[int] = Query(None, description="Quiz score (0-100) if action is quiz_completed"),
    flashcards_count: Optional[int] = Query(None, description="Number of flashcards viewed in the deck"),
    authorization: Optional[str] = Header(None)
):
    """
    Update user's progress on a topic.
    
    Actions:
    - 'start': Mark topic as started
    - 'flashcard_viewed': Set flashcard view count to the deck size
    - 'quiz_completed': Record quiz completion with score
    """
    try:
        user_id = await verify_user(authorization)
        
        # Validate topic exists (search both paths)
        topic = next((t for t in ALL_TOPICS if t["id"] == topic_id), None)
        if not topic:
            raise HTTPException(status_code=404, detail=f"Topic '{topic_id}' not found")
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        now = datetime.utcnow()
        
        if action == "start":
            # Insert or update - mark as started
            cursor.execute("""
                INSERT INTO user_topic_progress (user_id, topic_id, started_at, last_activity_at)
                VALUES (%s, %s, %s, %s)
                ON CONFLICT (user_id, topic_id) 
                DO UPDATE SET last_activity_at = EXCLUDED.last_activity_at
            """, (user_id, topic_id, now, now))
            
        elif action == "flashcard_viewed":
            # Set flashcard count to actual cards viewed (passed from frontend)
            # If not provided, use 10 as default (typical deck size)
            cards_viewed = flashcards_count or 10
            cursor.execute("""
                INSERT INTO user_topic_progress (user_id, topic_id, started_at, flashcards_viewed, last_activity_at)
                VALUES (%s, %s, %s, %s, %s)
                ON CONFLICT (user_id, topic_id) 
                DO UPDATE SET 
                    flashcards_viewed = GREATEST(user_topic_progress.flashcards_viewed, EXCLUDED.flashcards_viewed),
                    last_activity_at = EXCLUDED.last_activity_at
            """, (user_id, topic_id, now, cards_viewed, now))
            
        elif action == "quiz_completed":
            if quiz_score is None:
                raise HTTPException(status_code=400, detail="quiz_score required for quiz_completed action")
            
            # Mark as completed if score >= 70%, update best score
            is_passing = quiz_score >= 70
            
            cursor.execute("""
                INSERT INTO user_topic_progress (user_id, topic_id, started_at, completed_at, best_quiz_score, last_activity_at)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (user_id, topic_id) 
                DO UPDATE SET 
                    completed_at = CASE 
                        WHEN %s AND user_topic_progress.completed_at IS NULL THEN EXCLUDED.completed_at
                        ELSE user_topic_progress.completed_at 
                    END,
                    best_quiz_score = CASE 
                        WHEN EXCLUDED.best_quiz_score > COALESCE(user_topic_progress.best_quiz_score, 0) 
                        THEN EXCLUDED.best_quiz_score 
                        ELSE user_topic_progress.best_quiz_score 
                    END,
                    last_activity_at = EXCLUDED.last_activity_at
            """, (user_id, topic_id, now, now if is_passing else None, quiz_score, now, is_passing))
        else:
            raise HTTPException(status_code=400, detail=f"Invalid action: {action}")
        
        conn.commit()
        
        # Fetch updated progress
        cursor.execute("""
            SELECT topic_id, started_at, completed_at, best_quiz_score, 
                   flashcards_viewed, last_activity_at
            FROM user_topic_progress
            WHERE user_id = %s AND topic_id = %s
        """, (user_id, topic_id))
        
        row = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not row:
            raise HTTPException(status_code=500, detail="Failed to fetch updated progress")
        
        progress = {
            "topic_id": row[0],
            "started_at": row[1].isoformat() if row[1] else None,
            "completed_at": row[2].isoformat() if row[2] else None,
            "best_quiz_score": row[3],
            "flashcards_viewed": row[4] or 0,
            "last_activity_at": row[5].isoformat() if row[5] else None,
        }
        
        is_completed = progress["completed_at"] is not None
        
        # Get next topic if completed (check within same level first, then next level)
        next_topic = None
        if is_completed:
            # Check if in beginner path
            beginner_index = next((i for i, t in enumerate(BEGINNER_PATH) if t["id"] == topic_id), -1)
            if beginner_index >= 0:
                if beginner_index < len(BEGINNER_PATH) - 1:
                    # More beginner topics
                    next_topic = BEGINNER_PATH[beginner_index + 1]
                else:
                    # Last beginner topic - suggest first intermediate
                    next_topic = INTERMEDIATE_PATH[0] if INTERMEDIATE_PATH else None
            else:
                # Check if in intermediate path
                intermediate_index = next((i for i, t in enumerate(INTERMEDIATE_PATH) if t["id"] == topic_id), -1)
                if intermediate_index >= 0:
                    if intermediate_index < len(INTERMEDIATE_PATH) - 1:
                        next_topic = INTERMEDIATE_PATH[intermediate_index + 1]
                    else:
                        # Last intermediate topic - suggest first advanced
                        next_topic = ADVANCED_PATH[0] if ADVANCED_PATH else None
                else:
                    # Check if in advanced path
                    advanced_index = next((i for i, t in enumerate(ADVANCED_PATH) if t["id"] == topic_id), -1)
                    if advanced_index >= 0 and advanced_index < len(ADVANCED_PATH) - 1:
                        next_topic = ADVANCED_PATH[advanced_index + 1]
        
        logger.info(f"✅ Updated progress: user={user_id}, topic={topic_id}, action={action}")
        
        return {
            "topic_id": topic_id,
            "progress": progress,
            "is_completed": is_completed,
            "next_topic": next_topic,
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Update topic progress error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/learning/progress", tags=["Learning Path"])
async def get_all_progress(
    authorization: Optional[str] = Header(None)
):
    """
    Get user's progress on all topics in both learning paths.
    """
    try:
        user_id = await verify_user(authorization)
        
        db_url = os.getenv("DATABASE_URL")
        import psycopg2
        conn = psycopg2.connect(db_url)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT topic_id, started_at, completed_at, best_quiz_score, 
                   flashcards_viewed, last_activity_at
            FROM user_topic_progress
            WHERE user_id = %s
        """, (user_id,))
        
        progress_rows = cursor.fetchall()
        cursor.close()
        conn.close()
        
        # Build progress map
        progress_map = {}
        for row in progress_rows:
            progress_map[row[0]] = {
                "topic_id": row[0],
                "started_at": row[1].isoformat() if row[1] else None,
                "completed_at": row[2].isoformat() if row[2] else None,
                "best_quiz_score": row[3],
                "flashcards_viewed": row[4] or 0,
                "last_activity_at": row[5].isoformat() if row[5] else None,
            }
        
        # Build full response with all topics from both paths
        topics_with_progress = []
        for topic in ALL_TOPICS:
            topic_id = topic["id"]
            progress = progress_map.get(topic_id, {
                "topic_id": topic_id,
                "started_at": None,
                "completed_at": None,
                "best_quiz_score": None,
                "flashcards_viewed": 0,
                "last_activity_at": None,
            })
            topics_with_progress.append({
                "topic": topic,
                "progress": progress,
            })
        
        # Count completions per level
        beginner_completed = sum(
            1 for t in BEGINNER_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        intermediate_completed = sum(
            1 for t in INTERMEDIATE_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        advanced_completed = sum(
            1 for t in ADVANCED_PATH 
            if progress_map.get(t["id"], {}).get("completed_at")
        )
        total_completed = beginner_completed + intermediate_completed + advanced_completed
        
        return {
            "topics": topics_with_progress,
            "total_topics": len(ALL_TOPICS),
            "completed_topics": total_completed,
            "beginner_completed": beginner_completed,
            "intermediate_completed": intermediate_completed,
            "advanced_completed": advanced_completed,
            "beginner_total": len(BEGINNER_PATH),
            "intermediate_total": len(INTERMEDIATE_PATH),
            "advanced_total": len(ADVANCED_PATH),
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Get all progress error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    """Custom HTTP exception handler"""
    return JSONResponse(
        status_code=exc.status_code,
        content=ErrorResponse(
            error=exc.detail,
            detail=None
        ).model_dump()
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """General exception handler"""
    return JSONResponse(
        status_code=500,
        content=ErrorResponse(
            error="Internal server error",
            detail=str(exc)
        ).model_dump()
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)

